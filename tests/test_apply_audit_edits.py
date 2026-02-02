#!/usr/bin/env python3
"""
ABOUTME: Unit tests for apply_audit_edits.py
ABOUTME: Uses mock functions to construct various document content and rule scenarios
"""

import sys
import json
import tempfile
from pathlib import Path

# Add skills/doc-audit/scripts directory to path (must be before import)
_scripts_dir = Path(__file__).parent.parent / 'skills' / 'doc-audit' / 'scripts'
sys.path.insert(0, str(_scripts_dir))

import pytest  # noqa: E402
from lxml import etree  # noqa: E402
from unittest.mock import patch  # noqa: E402

import apply_audit_edits as apply_module  # noqa: E402  # type: ignore[import-not-found]
from apply_audit_edits import (  # noqa: E402  # type: ignore[import-not-found]
    AuditEditApplier, NS, DRAWING_PATTERN,
    strip_auto_numbering, EditItem, EditResult
)


# ============================================================
# XML Namespace Constants
# ============================================================

NSMAP = {
    'w': NS['w'],
    'w14': NS['w14'],
    'wp': NS['wp'],
}

# ============================================================
# Mock Helper Functions
# ============================================================

def create_paragraph_xml(text_content: str, para_id: str = "12345678") -> etree.Element:
    """
    Create a simple paragraph element with text.
    
    Args:
        text_content: Text to include in the paragraph
        para_id: w14:paraId attribute value
    
    Returns:
        lxml Element representing <w:p>
    """
    p = etree.Element(f'{{{NS["w"]}}}p', nsmap=NSMAP)
    p.set(f'{{{NS["w14"]}}}paraId', para_id)
    
    # Create run with text
    r = etree.SubElement(p, f'{{{NS["w"]}}}r')
    t = etree.SubElement(r, f'{{{NS["w"]}}}t')
    t.text = text_content
    
    return p


def create_paragraph_with_inline_image(
    text_before: str,
    img_id: str,
    img_name: str,
    text_after: str = "",
    para_id: str = "12345678"
) -> etree.Element:
    """
    Create a paragraph with text and inline image.
    
    Args:
        text_before: Text before the image
        img_id: Image id attribute
        img_name: Image name attribute
        text_after: Text after the image
        para_id: w14:paraId attribute value
    
    Returns:
        lxml Element representing <w:p>
    """
    p = etree.Element(f'{{{NS["w"]}}}p', nsmap=NSMAP)
    p.set(f'{{{NS["w14"]}}}paraId', para_id)
    
    # Run with text before image
    if text_before:
        r1 = etree.SubElement(p, f'{{{NS["w"]}}}r')
        t1 = etree.SubElement(r1, f'{{{NS["w"]}}}t')
        t1.text = text_before
    
    # Run with inline image
    r_img = etree.SubElement(p, f'{{{NS["w"]}}}r')
    drawing = etree.SubElement(r_img, f'{{{NS["w"]}}}drawing')
    inline = etree.SubElement(drawing, f'{{{NS["wp"]}}}inline')
    doc_pr = etree.SubElement(inline, f'{{{NS["wp"]}}}docPr')
    doc_pr.set('id', img_id)
    doc_pr.set('name', img_name)
    
    # Run with text after image
    if text_after:
        r2 = etree.SubElement(p, f'{{{NS["w"]}}}r')
        t2 = etree.SubElement(r2, f'{{{NS["w"]}}}t')
        t2.text = text_after
    
    return p


def create_paragraph_with_anchor_image(
    text_content: str,
    img_id: str,
    img_name: str,
    para_id: str = "12345678"
) -> etree.Element:
    """
    Create a paragraph with text and floating (anchor) image.
    Anchor images should be ignored by the system.
    
    Args:
        text_content: Text content
        img_id: Image id attribute
        img_name: Image name attribute
        para_id: w14:paraId attribute value
    
    Returns:
        lxml Element representing <w:p>
    """
    p = etree.Element(f'{{{NS["w"]}}}p', nsmap=NSMAP)
    p.set(f'{{{NS["w14"]}}}paraId', para_id)
    
    # Run with text
    r1 = etree.SubElement(p, f'{{{NS["w"]}}}r')
    t1 = etree.SubElement(r1, f'{{{NS["w"]}}}t')
    t1.text = text_content
    
    # Run with floating/anchor image (should be ignored)
    r_img = etree.SubElement(p, f'{{{NS["w"]}}}r')
    drawing = etree.SubElement(r_img, f'{{{NS["w"]}}}drawing')
    anchor = etree.SubElement(drawing, f'{{{NS["wp"]}}}anchor')  # anchor, not inline
    doc_pr = etree.SubElement(anchor, f'{{{NS["wp"]}}}docPr')
    doc_pr.set('id', img_id)
    doc_pr.set('name', img_name)
    
    return p


def create_paragraph_with_track_changes(
    text_before: str,
    deleted_text: str,
    inserted_text: str,
    text_after: str,
    para_id: str = "12345678"
) -> etree.Element:
    """
    Create a paragraph with track changes (w:del and w:ins).
    
    Args:
        text_before: Text before changes
        deleted_text: Text marked as deleted
        inserted_text: Text marked as inserted
        text_after: Text after changes
        para_id: w14:paraId attribute value
    
    Returns:
        lxml Element representing <w:p>
    """
    p = etree.Element(f'{{{NS["w"]}}}p', nsmap=NSMAP)
    p.set(f'{{{NS["w14"]}}}paraId', para_id)
    
    # Run with text before
    if text_before:
        r1 = etree.SubElement(p, f'{{{NS["w"]}}}r')
        t1 = etree.SubElement(r1, f'{{{NS["w"]}}}t')
        t1.text = text_before
    
    # Deleted text
    if deleted_text:
        del_elem = etree.SubElement(p, f'{{{NS["w"]}}}del')
        del_elem.set(f'{{{NS["w"]}}}id', "0")
        del_elem.set(f'{{{NS["w"]}}}author', "Test")
        r_del = etree.SubElement(del_elem, f'{{{NS["w"]}}}r')
        del_text = etree.SubElement(r_del, f'{{{NS["w"]}}}delText')
        del_text.text = deleted_text
    
    # Inserted text
    if inserted_text:
        ins_elem = etree.SubElement(p, f'{{{NS["w"]}}}ins')
        ins_elem.set(f'{{{NS["w"]}}}id', "1")
        ins_elem.set(f'{{{NS["w"]}}}author', "Test")
        r_ins = etree.SubElement(ins_elem, f'{{{NS["w"]}}}r')
        ins_t = etree.SubElement(r_ins, f'{{{NS["w"]}}}t')
        ins_t.text = inserted_text
    
    # Run with text after
    if text_after:
        r2 = etree.SubElement(p, f'{{{NS["w"]}}}r')
        t2 = etree.SubElement(r2, f'{{{NS["w"]}}}t')
        t2.text = text_after
    
    return p


def create_mock_applier():
    """
    Create a mock AuditEditApplier for testing internal methods.
    
    Returns:
        AuditEditApplier instance with mocked dependencies
    """
    with patch.object(AuditEditApplier, '_load_jsonl') as mock_load:
        mock_load.return_value = (
            {'source_file': '/tmp/test.docx', 'source_hash': 'sha256:abc123'},
            []
        )
        applier = AuditEditApplier.__new__(AuditEditApplier)
        applier.meta = {'source_file': '/tmp/test.docx', 'source_hash': 'sha256:abc123'}
        applier.edit_items = []
        applier.verbose = True
        applier.author = 'Test'
        applier.initials = 'Te'
        applier.next_change_id = 0
        applier.next_comment_id = 0
        applier.comments = []
        # Set a fixed timestamp for tests
        applier.operation_timestamp = '2025-01-01T00:00:00Z'
        return applier


def create_edit_item(
    uuid: str = "AAA",
    uuid_end: str = "BBB",
    violation_text: str = "Bad text",
    violation_reason: str = "Reason",
    fix_action: str = "replace",
    revised_text: str = "Good text",
    category: str = "semantic",
    rule_id: str = "R001"
) -> EditItem:
    """Create a minimal EditItem for CLI tests."""
    return EditItem(
        uuid=uuid,
        uuid_end=uuid_end,
        violation_text=violation_text,
        violation_reason=violation_reason,
        fix_action=fix_action,
        revised_text=revised_text,
        category=category,
        rule_id=rule_id
    )


def get_test_author(applier, category: str = "semantic") -> str:
    """Build test author using category suffix (matches production logic)."""
    return applier._author_for_item(create_edit_item(category=category))


# ============================================================
# Tests: _collect_runs_info_original
# ============================================================

class TestCollectRunsInfoOriginal:
    """Tests for _collect_runs_info_original method"""
    
    def test_simple_text(self):
        """Test paragraph with simple text"""
        applier = create_mock_applier()
        para = create_paragraph_xml("Hello World")
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        assert combined_text == "Hello World"
        assert len(runs_info) == 1
        assert runs_info[0]['text'] == "Hello World"
        assert runs_info[0]['start'] == 0
        assert runs_info[0]['end'] == 11
    
    def test_inline_image(self):
        """Test paragraph with inline image"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Before ",
            img_id="1",
            img_name="图片 1",
            text_after=" After"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        expected_img_str = '<drawing id="1" name="图片 1" />'
        assert expected_img_str in combined_text
        assert combined_text == f'Before {expected_img_str} After'
        
        # Find image run
        img_runs = [r for r in runs_info if r.get('is_drawing')]
        assert len(img_runs) == 1
        assert img_runs[0]['text'] == expected_img_str
    
    def test_anchor_image_ignored(self):
        """Test that floating (anchor) images are ignored"""
        applier = create_mock_applier()
        para = create_paragraph_with_anchor_image(
            text_content="Text with floating image",
            img_id="1",
            img_name="Float Image"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        # Should only contain text, no image placeholder
        assert combined_text == "Text with floating image"
        assert '<drawing' not in combined_text
        
        # No drawing runs
        img_runs = [r for r in runs_info if r.get('is_drawing')]
        assert len(img_runs) == 0
    
    def test_track_changes_deleted_text(self):
        """Test paragraph with deleted text (should be included in original)"""
        applier = create_mock_applier()
        para = create_paragraph_with_track_changes(
            text_before="Hello ",
            deleted_text="old",
            inserted_text="new",
            text_after=" World"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        # Original text should include deleted text, exclude inserted text
        assert combined_text == "Hello old World"
        assert "new" not in combined_text


# ============================================================
# Tests: _apply_delete
# ============================================================

class TestApplyDelete:
    """Tests for _apply_delete method"""
    
    def test_delete_simple_text(self):
        """Test deleting simple text"""
        applier = create_mock_applier()
        para = create_paragraph_xml("Hello World to delete")

        runs_info, _ = applier._collect_runs_info_original(para)

        result = applier._apply_delete(para, "to delete", "Test reason", runs_info, 12, get_test_author(applier))

        assert result == 'success'
        # Verify w:del element was created
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1
        # Verify comment was also created with -R suffix author
        assert len(applier.comments) == 1
        assert applier.comments[0]['text'] == "Test reason"
        assert applier.comments[0]['author'].endswith('-R')

    def test_delete_text_not_found(self):
        """Test deleting text that doesn't match position"""
        applier = create_mock_applier()
        para = create_paragraph_xml("Hello World")

        runs_info, _ = applier._collect_runs_info_original(para)

        # Position 50 is beyond text length
        result = applier._apply_delete(para, "missing", "Test reason", runs_info, 50, get_test_author(applier))

        assert result == 'fallback'


# ============================================================
# Tests: _apply_replace with Images
# ============================================================

class TestApplyReplaceWithImages:
    """Tests for _apply_replace method with image handling"""
    
    def test_replace_simple_text(self):
        """Test simple text replacement"""
        applier = create_mock_applier()
        para = create_paragraph_xml("Hello World")

        runs_info, combined_text = applier._collect_runs_info_original(para)

        # Simple text replacement
        violation_text = "Hello"
        revised_text = "Hi"

        match_start = combined_text.find(violation_text)
        result = applier._apply_replace(para, violation_text, revised_text, "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'
        # Verify w:del and w:ins elements were created
        del_elems = para.findall('.//w:del', NSMAP)
        ins_elems = para.findall('.//w:ins', NSMAP)
        assert len(del_elems) == 1
        assert len(ins_elems) == 1
        # Verify comment was also created with -R suffix author
        assert len(applier.comments) == 1
        assert applier.comments[0]['text'] == "Test reason"
        assert applier.comments[0]['author'].endswith('-R')

    def test_replace_insert_image_fallback(self):
        """Test that inserting images via replace triggers fallback"""
        applier = create_mock_applier()
        para = create_paragraph_xml("Hello World")

        runs_info, combined_text = applier._collect_runs_info_original(para)

        # Try to insert image in revised text
        violation_text = "Hello World"
        revised_text = 'Hello <drawing id="2" name="New Image" /> World'

        result = applier._apply_replace(para, violation_text, revised_text, "Test reason", runs_info, 0, get_test_author(applier))

        assert result == 'fallback'

    def test_replace_delete_image(self):
        """Test that deleting images via replace works"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Hello ",
            img_id="1",
            img_name="图片 1",
            text_after=" World"
        )

        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="图片 1" />'

        # Delete image from content
        violation_text = f"Hello {img_str} World"
        revised_text = "Hello World"

        match_start = combined_text.find(violation_text)
        result = applier._apply_replace(para, violation_text, revised_text, "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'


# ============================================================
# Tests: _apply_replace Equal Portions
# ============================================================

class TestApplyReplaceEqualPortions:
    """Tests for equal portion handling in _apply_replace"""
    
    def test_equal_portion_preserves_image(self):
        """Image in equal portion should be preserved (deepcopy)"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Hello ",
            img_id="1",
            img_name="图片 1",
            text_after=" World"
        )

        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="图片 1" />'

        # Replace "Hello" with "Hi", keep " <img> World" as equal
        violation_text = f"Hello {img_str} World"
        revised_text = f"Hi {img_str} World"

        match_start = combined_text.find(violation_text)
        result = applier._apply_replace(para, violation_text, revised_text, "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'

        # Verify image element still exists
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1

        # Verify inline docPr preserved
        inline = para.find('.//wp:inline', NSMAP)
        assert inline is not None
        doc_pr = inline.find('wp:docPr', NSMAP)
        assert doc_pr is not None
        assert doc_pr.get('id') == '1'
        assert doc_pr.get('name') == '图片 1'

    def test_equal_at_start_with_delete_at_end(self):
        """Equal portion at start, delete at end"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Keep ",
            img_id="1",
            img_name="图片 1",
            text_after=" Delete"
        )

        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="图片 1" />'

        # Delete " Delete" at end, keep "Keep <img>" as equal
        violation_text = f"Keep {img_str} Delete"
        revised_text = f"Keep {img_str}"

        match_start = combined_text.find(violation_text)
        result = applier._apply_replace(para, violation_text, revised_text, "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'

        # Verify image preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1

        # Verify w:del created for " Delete"
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1

    def test_equal_text_before_image_in_equal(self):
        """Equal portion has text before image, both should be preserved"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Delete ",
            img_id="1",
            img_name="图片 1",
            text_after=" End"
        )

        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="图片 1" />'

        # Delete "Delete", keep " <img> End" (space + image + text)
        violation_text = f"Delete {img_str} End"
        revised_text = f"{img_str} End"

        match_start = combined_text.find(violation_text)
        result = applier._apply_replace(para, violation_text, revised_text, "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'

        # Verify image preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1

        # Verify w:del created
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1


def create_paragraph_with_multiple_images(
    parts: list,
    para_id: str = "12345678"
) -> etree.Element:
    """
    Create a paragraph with multiple text/image parts.
    
    Args:
        parts: List of dicts, each with either 'text' or 'img' key
               e.g. [{'text': 'A '}, {'img': ('1', 'Img1')}, {'text': ' B'}]
        para_id: w14:paraId attribute value
    
    Returns:
        lxml Element representing <w:p>
    """
    p = etree.Element(f'{{{NS["w"]}}}p', nsmap=NSMAP)
    p.set(f'{{{NS["w14"]}}}paraId', para_id)
    
    for part in parts:
        if 'text' in part:
            r = etree.SubElement(p, f'{{{NS["w"]}}}r')
            t = etree.SubElement(r, f'{{{NS["w"]}}}t')
            t.text = part['text']
        elif 'img' in part:
            img_id, img_name = part['img']
            r_img = etree.SubElement(p, f'{{{NS["w"]}}}r')
            drawing = etree.SubElement(r_img, f'{{{NS["w"]}}}drawing')
            inline = etree.SubElement(drawing, f'{{{NS["wp"]}}}inline')
            doc_pr = etree.SubElement(inline, f'{{{NS["wp"]}}}docPr')
            doc_pr.set('id', img_id)
            doc_pr.set('name', img_name)
    
    return p


class TestApplyReplaceMultipleImages:
    """Tests for multiple images in equal portions"""

    def test_two_images_both_preserved(self):
        """Two images in equal portion should both be preserved"""
        applier = create_mock_applier()
        para = create_paragraph_with_multiple_images([
            {'text': 'A '},
            {'img': ('1', 'Img1')},
            {'text': ' B '},
            {'img': ('2', 'Img2')},
            {'text': ' C'}
        ])

        runs_info, combined_text = applier._collect_runs_info_original(para)
        img1_str = '<drawing id="1" name="Img1" />'
        img2_str = '<drawing id="2" name="Img2" />'

        # Replace "A" with "X", keep rest
        violation_text = f"A {img1_str} B {img2_str} C"
        revised_text = f"X {img1_str} B {img2_str} C"

        match_start = combined_text.find(violation_text)
        result = applier._apply_replace(para, violation_text, revised_text, "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'

        # Verify both images preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 2


# ============================================================
# Tests: _apply_manual (Comment)
# ============================================================

class TestApplyManual:
    """Tests for _apply_manual method"""
    
    def test_manual_comment_on_text(self):
        """Test adding comment to text"""
        applier = create_mock_applier()
        para = create_paragraph_xml("This is problematic text here")
        
        runs_info, _ = applier._collect_runs_info_original(para)
        
        result = applier._apply_manual(
            para, "problematic text",
            "This text is wrong", "Fix suggestion",
            runs_info, 8,
            get_test_author(applier)
        )
        
        assert result == 'success'
        # Verify comment elements
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1


# ============================================================
# Tests: _apply_delete with Images
# ============================================================

class TestApplyDeleteWithImages:
    """Tests for _apply_delete method with image content"""

    def test_delete_text_before_image(self):
        """Delete text appearing before an inline image, image should be preserved"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Delete this ",
            img_id="1",
            img_name="图片 1",
            text_after=" keep"
        )

        runs_info, combined_text = applier._collect_runs_info_original(para)

        # Delete "Delete this "
        result = applier._apply_delete(para, "Delete this ", "Test reason", runs_info, 0, get_test_author(applier))

        assert result == 'success'
        # Verify w:del element was created
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1
        # Verify image is preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1

    def test_delete_text_after_image(self):
        """Delete text appearing after an inline image, image should be preserved"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Keep ",
            img_id="1",
            img_name="图片 1",
            text_after=" delete this"
        )

        runs_info, combined_text = applier._collect_runs_info_original(para)

        # Find position of " delete this" after the image
        match_start = combined_text.find(" delete this")
        result = applier._apply_delete(para, " delete this", "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'
        # Verify w:del element was created
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1
        # Verify image is preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1

    def test_delete_image_placeholder(self):
        """Delete the image placeholder string itself"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Before ",
            img_id="1",
            img_name="图片 1",
            text_after=" After"
        )

        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="图片 1" />'

        # Delete the image placeholder
        match_start = combined_text.find(img_str)
        result = applier._apply_delete(para, img_str, "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'
        # Verify w:del element was created
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1

    def test_delete_text_with_image_in_middle(self):
        """Delete text in a paragraph where image is in the middle"""
        applier = create_mock_applier()
        para = create_paragraph_with_multiple_images([
            {'text': 'Start delete_me '},
            {'img': ('1', 'Img1')},
            {'text': ' End'}
        ])

        runs_info, combined_text = applier._collect_runs_info_original(para)

        # Delete "delete_me " before the image
        match_start = combined_text.find("delete_me ")
        result = applier._apply_delete(para, "delete_me ", "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'
        # Verify w:del element was created
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1
        # Verify image is still present
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1

    def test_delete_text_between_two_images(self):
        """Delete text between two images"""
        applier = create_mock_applier()
        para = create_paragraph_with_multiple_images([
            {'img': ('1', 'Img1')},
            {'text': ' delete me '},
            {'img': ('2', 'Img2')}
        ])

        runs_info, combined_text = applier._collect_runs_info_original(para)

        # Delete " delete me " between the images
        match_start = combined_text.find(" delete me ")
        result = applier._apply_delete(para, " delete me ", "Test reason", runs_info, match_start, get_test_author(applier))

        assert result == 'success'
        # Verify w:del element was created
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1
        # Verify both images are preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 2


# ============================================================
# Tests: _apply_manual with Images
# ============================================================

class TestApplyManualWithImages:
    """Tests for _apply_manual method with image content"""
    
    def test_manual_comment_on_text_before_image(self):
        """Add comment to text before an inline image"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Mark this text ",
            img_id="1",
            img_name="图片 1",
            text_after=" after"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        result = applier._apply_manual(
            para, "Mark this",
            "This text needs review", "Suggestion here",
            runs_info, 0,
            get_test_author(applier)
        )
        
        assert result == 'success'
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1
        # Verify image is preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1
        # Verify comment was recorded
        assert len(applier.comments) == 1
    
    def test_manual_comment_on_text_after_image(self):
        """Add comment to text after an inline image"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Before ",
            img_id="1",
            img_name="图片 1",
            text_after=" mark this text"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        # Find position of "mark this" after the image
        match_start = combined_text.find("mark this")
        result = applier._apply_manual(
            para, "mark this",
            "This text needs fixing", "Fix suggestion",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1
        # Verify image is preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1
    
    def test_manual_comment_on_image_placeholder(self):
        """Add comment directly on the image placeholder string"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Before ",
            img_id="1",
            img_name="图片 1",
            text_after=" After"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="图片 1" />'
        
        # Add comment to the image placeholder
        match_start = combined_text.find(img_str)
        result = applier._apply_manual(
            para, img_str,
            "This image needs replacement", "Use a better image",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1
        # Verify comment was recorded
        assert len(applier.comments) == 1
        assert "This image needs replacement" in applier.comments[0]['text']
    
    def test_manual_comment_spanning_text_and_image(self):
        """Add comment to text that includes an image placeholder"""
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="See ",
            img_id="1",
            img_name="图片 1",
            text_after=" above"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="图片 1" />'
        
        # Add comment spanning "See <image> above"
        violation_text = f"See {img_str} above"
        match_start = combined_text.find(violation_text)
        result = applier._apply_manual(
            para, violation_text,
            "This section needs work", "Improve the description",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1
    
    def test_manual_comment_with_multiple_images(self):
        """Add comment in a paragraph with multiple images"""
        applier = create_mock_applier()
        para = create_paragraph_with_multiple_images([
            {'text': 'Text A '},
            {'img': ('1', 'Img1')},
            {'text': ' mark this '},
            {'img': ('2', 'Img2')},
            {'text': ' Text B'}
        ])
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        # Add comment to "mark this" between the images
        match_start = combined_text.find("mark this")
        result = applier._apply_manual(
            para, "mark this",
            "This needs attention", "Suggestion",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1
        # Verify both images are preserved
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 2
        # Verify comment was recorded
        assert len(applier.comments) == 1
    
    def test_manual_comment_on_image_between_text(self):
        """Add comment on image that is between text runs"""
        applier = create_mock_applier()
        para = create_paragraph_with_multiple_images([
            {'text': 'Start '},
            {'img': ('1', 'ProblemImage')},
            {'text': ' End'}
        ])
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="ProblemImage" />'
        
        # Add comment to the image
        match_start = combined_text.find(img_str)
        result = applier._apply_manual(
            para, img_str,
            "Image quality is low", "Replace with high-res version",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        # Verify comment was recorded with correct content
        assert len(applier.comments) == 1
        assert "Image quality is low" in applier.comments[0]['text']
        assert "Replace with high-res version" in applier.comments[0]['text']


# ============================================================
# Tests: DRAWING_PATTERN regex
# ============================================================

class TestDrawingPattern:
    """Tests for DRAWING_PATTERN regex"""
    
    def test_matches_valid_drawing(self):
        """Test pattern matches valid drawing placeholder"""
        valid_cases = [
            '<drawing id="1" name="图片 1" />',
            '<drawing id="123" name="Image" />',
            '<drawing id="0" name="" />',
        ]
        for case in valid_cases:
            assert DRAWING_PATTERN.fullmatch(case), f"Should match: {case}"
    
    def test_no_match_invalid(self):
        """Test pattern doesn't match invalid formats"""
        invalid_cases = [
            '<drawing id="1" name="图片 1">content</drawing>',  # has content
            '<draw id="1" name="test" />',  # wrong tag
            'text <drawing id="1" name="Test" /> more',  # mixed with text
        ]
        for case in invalid_cases:
            assert not DRAWING_PATTERN.fullmatch(case), f"Should not fullmatch: {case}"
    
    def test_search_finds_in_text(self):
        """Test search finds drawing in mixed text"""
        text = 'Hello <drawing id="1" name="图片 1" /> World'
        assert DRAWING_PATTERN.search(text)


# ============================================================
# Tests: Auto-numbering stripping
# ============================================================

class TestStripAutoNumbering:
    """Tests for strip_auto_numbering helper"""
    
    def test_strip_numeric(self):
        """Test stripping numeric prefixes"""
        assert strip_auto_numbering("1. Introduction") == ("Introduction", True)
        assert strip_auto_numbering("1.1 Details") == ("Details", True)
        assert strip_auto_numbering("1) First") == ("First", True)
    
    def test_strip_alphabetic(self):
        """Test stripping alphabetic prefixes"""
        assert strip_auto_numbering("a. First item") == ("First item", True)
        assert strip_auto_numbering("A) Section") == ("Section", True)
    
    def test_strip_bullet(self):
        """Test stripping bullet"""
        assert strip_auto_numbering("• Note") == ("Note", True)
    
    def test_no_strip_normal(self):
        """Test no stripping for normal text"""
        assert strip_auto_numbering("Normal text") == ("Normal text", False)


# ============================================================
# Tests: _find_revision_ancestor
# ============================================================

class TestFindRevisionAncestor:
    """Tests for _find_revision_ancestor helper method"""
    
    def test_find_revision_ancestor_normal_run(self):
        """Normal run (not in revision) should return None"""
        applier = create_mock_applier()
        para = create_paragraph_xml("Hello World")
        
        # Get the run element
        run = para.find('.//w:r', NSMAP)
        
        result = applier._find_revision_ancestor(run, para)
        assert result is None
    
    def test_find_revision_ancestor_inside_del(self):
        """Run inside w:del should return the del element"""
        applier = create_mock_applier()
        para = create_paragraph_with_track_changes(
            text_before="Before ",
            deleted_text="deleted",
            inserted_text="",
            text_after=" after"
        )
        
        # Find the w:del element and its contained run
        del_elem = para.find('.//w:del', NSMAP)
        run_in_del = del_elem.find('.//w:r', NSMAP)
        
        result = applier._find_revision_ancestor(run_in_del, para)
        assert result is not None
        assert result.tag == f'{{{NS["w"]}}}del'
        assert result is del_elem
    
    def test_find_revision_ancestor_inside_ins(self):
        """Run inside w:ins should return the ins element"""
        applier = create_mock_applier()
        para = create_paragraph_with_track_changes(
            text_before="Before ",
            deleted_text="",
            inserted_text="inserted",
            text_after=" after"
        )
        
        # Find the w:ins element and its contained run
        ins_elem = para.find('.//w:ins', NSMAP)
        run_in_ins = ins_elem.find('.//w:r', NSMAP)
        
        result = applier._find_revision_ancestor(run_in_ins, para)
        assert result is not None
        assert result.tag == f'{{{NS["w"]}}}ins'
        assert result is ins_elem


# ============================================================
# Tests: _apply_manual with Revision Content (New Tests)
# ============================================================

class TestApplyManualWithRevisions:
    """Tests for _apply_manual method preserving revision structure"""
    
    def test_manual_comment_on_deleted_text(self):
        """
        Add comment to text that was deleted by a previous rule.
        Comment markers should be inserted outside the w:del container.
        """
        applier = create_mock_applier()
        para = create_paragraph_with_track_changes(
            text_before="Before ",
            deleted_text="mark this deleted text",
            inserted_text="replacement",
            text_after=" after"
        )
        
        # Collect original text (includes deleted text, excludes inserted)
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        # Verify original text includes the deleted content
        assert "mark this deleted text" in combined_text
        assert "replacement" not in combined_text
        
        # Find position of "mark this" in the deleted text
        match_start = combined_text.find("mark this deleted text")
        
        result = applier._apply_manual(
            para, "mark this deleted text",
            "This deleted text needs review", "Suggestion",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1
        
        # Verify w:del element is still present and intact
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1
        
        # Verify commentRangeStart is BEFORE w:del (sibling, not inside)
        del_elem = del_elems[0]
        del_idx = list(para).index(del_elem)
        range_start_elem = range_start[0]
        range_start_idx = list(para).index(range_start_elem)
        assert range_start_idx < del_idx, "commentRangeStart should be before w:del"
        
        # Verify comment was recorded
        assert len(applier.comments) == 1
    
    def test_manual_comment_spanning_normal_and_deleted(self):
        """
        Add comment spanning normal text and deleted text.
        Start marker should be in normal run, end marker after w:del container.
        """
        applier = create_mock_applier()
        para = create_paragraph_with_track_changes(
            text_before="mark this normal ",
            deleted_text="and deleted",
            inserted_text="",
            text_after=" after"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        # Comment spans "this normal and deleted"
        violation_text = "this normal and deleted"
        match_start = combined_text.find(violation_text)
        
        result = applier._apply_manual(
            para, violation_text,
            "This span needs review", "Fix suggestion",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1
        
        # Verify w:del element is still present
        del_elems = para.findall('.//w:del', NSMAP)
        assert len(del_elems) == 1
        
        # Verify comment was recorded
        assert len(applier.comments) == 1
    
    def test_manual_preserves_multiple_runs_in_range(self):
        """
        Add comment to text spanning multiple runs.
        All runs should be preserved, not replaced.
        """
        applier = create_mock_applier()
        # Create paragraph with multiple runs
        para = create_paragraph_with_multiple_images([
            {'text': 'Run1 '},
            {'text': 'Run2 '},
            {'text': 'Run3'}
        ])
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        
        # Count runs before
        runs_before = len(para.findall('.//w:r', NSMAP))
        
        # Comment spans "Run2"
        match_start = combined_text.find("Run2")
        result = applier._apply_manual(
            para, "Run2",
            "Middle run needs attention", "Suggestion",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1
        
        # Verify runs are preserved (may have additional runs for split, but original structure intact)
        runs_after = para.findall('.//w:r', NSMAP)
        # Should have at least the original runs (may have more due to splitting)
        assert len(runs_after) >= runs_before
        
        # Verify comment was recorded
        assert len(applier.comments) == 1
    
    def test_manual_with_image_in_range_preserves_image(self):
        """
        Add comment to text that includes an image.
        Image run should be preserved, not replaced with text.
        """
        applier = create_mock_applier()
        para = create_paragraph_with_inline_image(
            text_before="Before ",
            img_id="1",
            img_name="TestImg",
            text_after=" After"
        )
        
        runs_info, combined_text = applier._collect_runs_info_original(para)
        img_str = '<drawing id="1" name="TestImg" />'
        
        # Verify image is in original text
        assert img_str in combined_text
        
        # Comment spans text including image
        violation_text = f"Before {img_str} After"
        match_start = combined_text.find(violation_text)
        
        result = applier._apply_manual(
            para, violation_text,
            "This section includes image", "Suggestion",
            runs_info, match_start,
            get_test_author(applier)
        )
        
        assert result == 'success'
        
        # Verify image element is still present
        img_elems = para.findall('.//w:drawing', NSMAP)
        assert len(img_elems) == 1
        
        # Verify inline structure preserved
        inline = para.find('.//wp:inline', NSMAP)
        assert inline is not None
        doc_pr = inline.find('wp:docPr', NSMAP)
        assert doc_pr is not None
        assert doc_pr.get('id') == '1'
        assert doc_pr.get('name') == 'TestImg'
        
        # Verify comment elements were created
        range_start = para.findall('.//w:commentRangeStart', NSMAP)
        range_end = para.findall('.//w:commentRangeEnd', NSMAP)
        assert len(range_start) == 1
        assert len(range_end) == 1


# ============================================================
# Tests: _load_jsonl strict uuid_end validation
# ============================================================

class TestLoadJsonlStrictValidation:
    """Tests for _load_jsonl strict uuid_end validation"""
    
    def test_load_jsonl_with_uuid_end_flat_format(self):
        """Test loading flat format JSONL with uuid_end field"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.jsonl', delete=False) as f:
            # Write meta line
            meta = {'type': 'meta', 'source_file': '/tmp/test.docx', 'source_hash': 'sha256:abc123'}
            json.dump(meta, f)
            f.write('\n')
            
            # Write edit item with uuid_end
            item = {
                'uuid': 'AAAAAAAA',
                'uuid_end': 'BBBBBBBB',
                'violation_text': 'bad text',
                'violation_reason': 'wrong',
                'fix_action': 'delete',
                'revised_text': '',
                'category': 'test',
                'rule_id': 'R001'
            }
            json.dump(item, f)
            f.write('\n')
            
            f.flush()
            
            # Create applier with mocked dependencies
            with patch.object(AuditEditApplier, '__init__', lambda x, *args, **kwargs: None):
                applier = AuditEditApplier.__new__(AuditEditApplier)
                applier.jsonl_path = Path(f.name)
                
                meta_loaded, items_loaded = applier._load_jsonl()
                
                assert len(items_loaded) == 1
                assert items_loaded[0].uuid == 'AAAAAAAA'
                assert items_loaded[0].uuid_end == 'BBBBBBBB'
    
    def test_load_jsonl_missing_uuid_end_flat_format_raises(self):
        """Test that missing uuid_end in flat format raises ValueError"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.jsonl', delete=False) as f:
            # Write meta line
            meta = {'type': 'meta', 'source_file': '/tmp/test.docx', 'source_hash': 'sha256:abc123'}
            json.dump(meta, f)
            f.write('\n')
            
            # Write edit item WITHOUT uuid_end
            item = {
                'uuid': 'AAAAAAAA',
                # 'uuid_end' is missing
                'violation_text': 'bad text',
                'violation_reason': 'wrong',
                'fix_action': 'delete',
                'revised_text': '',
                'category': 'test',
                'rule_id': 'R001'
            }
            json.dump(item, f)
            f.write('\n')
            
            f.flush()
            
            # Create applier with mocked dependencies
            with patch.object(AuditEditApplier, '__init__', lambda x, *args, **kwargs: None):
                applier = AuditEditApplier.__new__(AuditEditApplier)
                applier.jsonl_path = Path(f.name)
                
                with pytest.raises(ValueError, match="Missing 'uuid_end' field"):
                    applier._load_jsonl()
    
    def test_load_jsonl_with_uuid_end_nested_format(self):
        """Test loading nested format JSONL (from run_audit.py) with uuid_end"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.jsonl', delete=False) as f:
            # Write meta line
            meta = {'type': 'meta', 'source_file': '/tmp/test.docx', 'source_hash': 'sha256:abc123'}
            json.dump(meta, f)
            f.write('\n')
            
            # Write paragraph with violations (nested format)
            para = {
                'uuid': 'AAAAAAAA',
                'uuid_end': 'BBBBBBBB',  # Block-level uuid_end
                'p_heading': 'Section 1',
                'p_content': 'Content text',
                'violations': [
                    {
                        'violation_text': 'bad text',
                        'violation_reason': 'wrong',
                        'fix_action': 'replace',
                        'revised_text': 'good text',
                        'category': 'test',
                        'rule_id': 'R001'
                        # uuid_end inherited from paragraph level
                    }
                ]
            }
            json.dump(para, f)
            f.write('\n')
            
            f.flush()
            
            # Create applier with mocked dependencies
            with patch.object(AuditEditApplier, '__init__', lambda x, *args, **kwargs: None):
                applier = AuditEditApplier.__new__(AuditEditApplier)
                applier.jsonl_path = Path(f.name)
                
                meta_loaded, items_loaded = applier._load_jsonl()
                
                assert len(items_loaded) == 1
                assert items_loaded[0].uuid == 'AAAAAAAA'
                assert items_loaded[0].uuid_end == 'BBBBBBBB'
                assert items_loaded[0].heading == 'Section 1'
    
    def test_load_jsonl_missing_uuid_end_nested_format_raises(self):
        """Test that missing uuid_end in nested format raises ValueError"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.jsonl', delete=False) as f:
            # Write meta line
            meta = {'type': 'meta', 'source_file': '/tmp/test.docx', 'source_hash': 'sha256:abc123'}
            json.dump(meta, f)
            f.write('\n')
            
            # Write paragraph with violations WITHOUT uuid_end
            para = {
                'uuid': 'AAAAAAAA',
                # 'uuid_end' is missing at both levels
                'p_heading': 'Section 1',
                'p_content': 'Content text',
                'violations': [
                    {
                        'violation_text': 'bad text',
                        'violation_reason': 'wrong',
                        'fix_action': 'replace',
                        'revised_text': 'good text',
                        'category': 'test',
                        'rule_id': 'R001'
                    }
                ]
            }
            json.dump(para, f)
            f.write('\n')
            
            f.flush()
            
            # Create applier with mocked dependencies
            with patch.object(AuditEditApplier, '__init__', lambda x, *args, **kwargs: None):
                applier = AuditEditApplier.__new__(AuditEditApplier)
                applier.jsonl_path = Path(f.name)
                
                with pytest.raises(ValueError, match="Missing 'uuid_end' field"):
                    applier._load_jsonl()


# ============================================================
# Tests: _iter_paragraphs_in_range
# ============================================================

def create_mock_body_with_paragraphs(para_ids: list) -> etree.Element:
    """
    Create a mock document body with multiple paragraphs.
    
    Args:
        para_ids: List of paraId values for each paragraph
    
    Returns:
        lxml Element representing document body
    """
    body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
    
    for para_id in para_ids:
        p = etree.SubElement(body, f'{{{NS["w"]}}}p')
        p.set(f'{{{NS["w14"]}}}paraId', para_id)
        r = etree.SubElement(p, f'{{{NS["w"]}}}r')
        t = etree.SubElement(r, f'{{{NS["w"]}}}t')
        t.text = f"Paragraph {para_id}"
    
    return body


class TestIterParagraphsInRange:
    """Tests for _iter_paragraphs_in_range method"""
    
    def test_iter_single_paragraph(self):
        """Test iteration over a single paragraph (uuid == uuid_end)"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB', 'CCC'])
        applier.body_elem = body
        
        # Find start paragraph
        start_para = body.find(f'.//w:p[@w14:paraId="BBB"]', NSMAP)
        
        # Iterate with same uuid_end (single paragraph)
        paras = list(applier._iter_paragraphs_in_range(start_para, 'BBB'))
        
        assert len(paras) == 1
        assert paras[0].get(f'{{{NS["w14"]}}}paraId') == 'BBB'
    
    def test_iter_multiple_paragraphs(self):
        """Test iteration over multiple paragraphs"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB', 'CCC', 'DDD', 'EEE'])
        applier.body_elem = body
        
        # Find start paragraph
        start_para = body.find(f'.//w:p[@w14:paraId="BBB"]', NSMAP)
        
        # Iterate from BBB to DDD
        paras = list(applier._iter_paragraphs_in_range(start_para, 'DDD'))
        
        assert len(paras) == 3
        para_ids = [p.get(f'{{{NS["w14"]}}}paraId') for p in paras]
        assert para_ids == ['BBB', 'CCC', 'DDD']
    
    def test_iter_stops_at_uuid_end(self):
        """Test that iteration stops at uuid_end (inclusive)"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB', 'CCC', 'DDD', 'EEE'])
        applier.body_elem = body
        
        # Find start paragraph
        start_para = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)
        
        # Iterate from AAA to CCC - should not include DDD, EEE
        paras = list(applier._iter_paragraphs_in_range(start_para, 'CCC'))
        
        assert len(paras) == 3
        para_ids = [p.get(f'{{{NS["w14"]}}}paraId') for p in paras]
        assert para_ids == ['AAA', 'BBB', 'CCC']
        assert 'DDD' not in para_ids
        assert 'EEE' not in para_ids
    
    def test_iter_with_nonexistent_uuid_end(self):
        """Test iteration when uuid_end doesn't exist (iterates to end)"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB', 'CCC'])
        applier.body_elem = body
        
        # Find start paragraph
        start_para = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)
        
        # uuid_end doesn't exist - should iterate all remaining paragraphs
        paras = list(applier._iter_paragraphs_in_range(start_para, 'ZZZZZ'))
        
        assert len(paras) == 3
        para_ids = [p.get(f'{{{NS["w14"]}}}paraId') for p in paras]
        assert para_ids == ['AAA', 'BBB', 'CCC']
    
    def test_iter_start_not_in_body(self):
        """Test iteration when start_node is not in body"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB', 'CCC'])
        applier.body_elem = body
        
        # Create a detached paragraph (not in body)
        detached_para = create_paragraph_xml("Detached", "XXX")
        
        # Should return empty list
        paras = list(applier._iter_paragraphs_in_range(detached_para, 'BBB'))
        
        assert len(paras) == 0


# ============================================================
# Tests: main exit code behavior
# ============================================================

class TestMainExitCodeBehavior:
    """Tests for main() exit codes on warnings/failures."""
    
    class DummyApplier:
        results = []
        
        def __init__(self, jsonl_file, output_path=None, author=None, initials=None, skip_hash=False, verbose=False):
            self.source_path = Path(jsonl_file)
            self.output_path = Path(output_path) if output_path else Path("out.docx")
            self.edit_items = [1, 2]
        
        def apply(self):
            return self.__class__.results
        
        def save(self, dry_run=False):
            return None
        
        def save_failed_items(self):
            return None
    
    def test_main_returns_zero_with_failures(self, monkeypatch):
        """main() should return 0 even when warnings or failures exist."""
        item = create_edit_item()
        self.DummyApplier.results = [
            EditResult(True, item, "fallback", warning=True),
            EditResult(False, item, "failed"),
        ]
        
        monkeypatch.setattr(apply_module, "AuditEditApplier", self.DummyApplier)
        monkeypatch.setattr(sys, "argv", ["apply_audit_edits.py", "input.jsonl", "--dry-run"])
        
        exit_code = apply_module.main()
        assert exit_code == 0


# ============================================================
# Helper: Create table cell with multiple paragraphs
# ============================================================

def create_table_cell_with_paragraphs(para_contents: list, para_ids: list) -> etree.Element:
    """
    Create a table cell (<w:tc>) with multiple paragraphs.
    This simulates the XML structure of a Word table cell with multiple paragraphs.
    
    Args:
        para_contents: List of text content for each paragraph
        para_ids: List of paraId values for each paragraph
    
    Returns:
        lxml Element representing <w:tc>
    """
    tc = etree.Element(f'{{{NS["w"]}}}tc', nsmap=NSMAP)
    
    for content, para_id in zip(para_contents, para_ids):
        p = etree.SubElement(tc, f'{{{NS["w"]}}}p')
        p.set(f'{{{NS["w14"]}}}paraId', para_id)
        
        r = etree.SubElement(p, f'{{{NS["w"]}}}r')
        t = etree.SubElement(r, f'{{{NS["w"]}}}t')
        t.text = content
    
    return tc


# ============================================================
# Tests: _collect_runs_info_across_paragraphs
# ============================================================

class TestCollectRunsInfoAcrossParagraphs:
    """Tests for _collect_runs_info_across_paragraphs method"""

    def test_single_paragraph_returns_correct_flag(self):
        """Single paragraph should return is_cross_paragraph=False"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA'])
        applier.body_elem = body

        start_para = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(start_para, 'AAA')

        assert boundary_error is None
        assert is_cross_paragraph is False
        assert 'Paragraph AAA' in combined_text

    def test_multiple_paragraphs_returns_correct_flag(self):
        """Multiple paragraphs should return is_cross_paragraph=True"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB', 'CCC'])
        applier.body_elem = body

        start_para = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(start_para, 'CCC')

        assert boundary_error is None
        assert is_cross_paragraph is True
        assert 'Paragraph AAA' in combined_text
        assert 'Paragraph CCC' in combined_text

    def test_paragraph_boundaries_as_newlines(self):
        """Paragraph boundaries should be converted to \\n"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB'])
        applier.body_elem = body

        start_para = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(start_para, 'BBB')

        assert boundary_error is None
        # Should have newline between paragraphs
        assert '\n' in combined_text
        assert combined_text == 'Paragraph AAA\nParagraph BBB'

    def test_runs_contain_para_elem_reference(self):
        """Runs should contain para_elem reference"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB'])
        applier.body_elem = body

        start_para = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(start_para, 'BBB')

        assert boundary_error is None
        # Filter out boundary markers
        real_runs = [r for r in runs_info if not r.get('is_para_boundary', False)]

        # All real runs should have para_elem
        for run in real_runs:
            assert 'para_elem' in run
            assert run['para_elem'] is not None

    def test_para_boundary_markers_flagged(self):
        """Paragraph boundary runs should be flagged with is_para_boundary=True"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB', 'CCC'])
        applier.body_elem = body

        start_para = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(start_para, 'CCC')

        assert boundary_error is None
        # Find boundary markers
        boundary_runs = [r for r in runs_info if r.get('is_para_boundary', False)]

        # Should have 2 boundary markers (AAA->BBB, BBB->CCC)
        assert len(boundary_runs) == 2

        # Boundary markers should all be '\n'
        for run in boundary_runs:
            assert run['text'] == '\n'

    def test_no_boundary_after_last_paragraph(self):
        """No boundary marker should be added after the last paragraph (uuid_end)"""
        applier = create_mock_applier()
        body = create_mock_body_with_paragraphs(['AAA', 'BBB'])
        applier.body_elem = body

        start_para = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(start_para, 'BBB')

        assert boundary_error is None
        # Should have exactly 1 boundary marker (AAA->BBB)
        boundary_runs = [r for r in runs_info if r.get('is_para_boundary', False)]
        assert len(boundary_runs) == 1


# ============================================================
# Tests: Cross-paragraph search
# ============================================================

class TestCrossParagraphSearch:
    """Tests for cross-paragraph text search in _process_item"""
    
    def test_single_para_match_found_first(self):
        """Single paragraph match should be found first (no cross-para search)"""
        applier = create_mock_applier()
        
        # Create body with multiple paragraphs
        body = create_mock_body_with_paragraphs(['AAA', 'BBB', 'CCC'])
        applier.body_elem = body
        
        # Modify BBB to have specific content
        para_bbb = body.find(f'.//w:p[@w14:paraId="BBB"]', NSMAP)
        para_bbb.find('.//w:t', NSMAP).text = "This is a violation"
        
        # Create edit item
        item = create_edit_item(
            uuid='AAA',
            uuid_end='CCC',
            violation_text='This is a violation',
            fix_action='manual'
        )
        
        # Process item (should find in single paragraph)
        result = applier._process_item(item)
        
        # Should succeed without using cross-paragraph mode
        assert result.success is True


# ============================================================
# Tests: Cross-paragraph manual comments
# ============================================================

class TestApplyManualCrossParagraph:
    """Tests for _apply_manual with cross-paragraph content"""
    
    def test_manual_comment_across_paragraphs(self):
        """Cross-paragraph comment should correctly insert commentRangeStart/End"""
        applier = create_mock_applier()
        
        # Create table cell with multiple paragraphs
        tc = create_table_cell_with_paragraphs(
            ['Line 1', 'Line 2', 'Line 3'],
            ['AAA', 'BBB', 'CCC']
        )
        
        # Create mock body and add cell paragraphs to it
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        for p in tc.findall('.//w:p', NSMAP):
            body.append(p)
        applier.body_elem = body
        
        # Get paragraphs
        para_aaa = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        # Collect runs across paragraphs
        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(para_aaa, 'CCC')

        assert boundary_error is None
        # Verify combined text
        assert combined_text == 'Line 1\nLine 2\nLine 3'
        assert is_cross_paragraph is True

        # Apply manual comment to text spanning all 3 paragraphs
        violation_text = 'Line 1\nLine 2\nLine 3'
        match_start = 0
        
        result = applier._apply_manual(
            para_aaa, violation_text,
            "Cross-paragraph issue", "Fix suggestion",
            runs_info, match_start,
            get_test_author(applier),
            is_cross_paragraph=True
        )
        
        assert result == 'success'
        
        # Verify comment markers were created
        all_range_starts = body.findall('.//w:commentRangeStart', NSMAP)
        all_range_ends = body.findall('.//w:commentRangeEnd', NSMAP)
        
        assert len(all_range_starts) == 1
        assert len(all_range_ends) == 1
        
        # Verify comment was recorded
        assert len(applier.comments) == 1
    
    def test_manual_filters_para_boundary_markers(self):
        """Should filter paragraph boundary markers, only process real runs"""
        applier = create_mock_applier()
        
        # Create paragraphs
        body = create_mock_body_with_paragraphs(['AAA', 'BBB'])
        applier.body_elem = body
        
        para_aaa = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        # Collect runs (will include boundary marker)
        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(para_aaa, 'BBB')

        assert boundary_error is None
        # Verify boundary marker exists
        boundary_runs = [r for r in runs_info if r.get('is_para_boundary', False)]
        assert len(boundary_runs) == 1

        # Apply manual - should filter out boundary marker
        violation_text = combined_text  # Full text including \n
        match_start = 0
        
        result = applier._apply_manual(
            para_aaa, violation_text,
            "Test", "Test",
            runs_info, match_start,
            get_test_author(applier),
            is_cross_paragraph=True
        )
        
        # Should succeed (boundary marker filtered)
        assert result == 'success'


# ============================================================
# Tests: Cross-paragraph fallback for delete/replace
# ============================================================

class TestCrossParagraphFallback:
    """Tests for cross-paragraph fallback behavior in delete/replace"""
    
    def test_delete_with_single_para_content_proceeds(self):
        """Search cross-para but match in single para → normal delete"""
        applier = create_mock_applier()
        
        # Create multiple paragraphs, but violation only in first
        body = create_mock_body_with_paragraphs(['AAA', 'BBB'])
        applier.body_elem = body
        
        # Modify AAA to have violation text
        para_aaa = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)
        para_aaa.find('.//w:t', NSMAP).text = "Delete this text"

        # Collect runs across paragraphs
        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(para_aaa, 'BBB')

        assert boundary_error is None
        assert is_cross_paragraph is True  # Search is cross-para
        
        # But actual match is only in AAA
        violation_text = "Delete this"
        match_start = combined_text.find(violation_text)
        assert match_start == 0
        
        # Find affected runs
        match_end = match_start + len(violation_text)
        affected = applier._find_affected_runs(runs_info, match_start, match_end)
        real_runs = [r for r in affected if not r.get('is_para_boundary', False)]
        
        # Check actual paragraphs spanned
        para_elems = set(r.get('para_elem') for r in real_runs if r.get('para_elem') is not None)
        
        # Should only span 1 paragraph (AAA)
        assert len(para_elems) == 1
        
        # Delete should proceed (not fallback)
        # This simulates the logic in _process_item for delete operation
        if len(para_elems) == 1:
            target_para = real_runs[0].get('para_elem')
            result = applier._apply_delete(
                target_para, violation_text,
                "Test reason",
                runs_info, match_start,
                get_test_author(applier)
            )
            assert result == 'success'

    def test_replace_with_single_para_content_proceeds(self):
        """Search cross-para but match in single para → normal replace"""
        applier = create_mock_applier()

        # Create multiple paragraphs
        body = create_mock_body_with_paragraphs(['AAA', 'BBB'])
        applier.body_elem = body

        # Modify BBB to have violation text
        para_bbb = body.find(f'.//w:p[@w14:paraId="BBB"]', NSMAP)
        para_bbb.find('.//w:t', NSMAP).text = "Bad text here"

        para_aaa = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        # Collect runs across paragraphs
        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(para_aaa, 'BBB')

        assert boundary_error is None
        assert is_cross_paragraph is True

        # Match is only in BBB (after \n)
        violation_text = "Bad text"
        match_start = combined_text.find(violation_text)

        # Find affected runs
        match_end = match_start + len(violation_text)
        affected = applier._find_affected_runs(runs_info, match_start, match_end)
        real_runs = [r for r in affected if not r.get('is_para_boundary', False)]

        # Check actual paragraphs spanned
        para_elems = set(r.get('para_elem') for r in real_runs if r.get('para_elem') is not None)

        # Should only span 1 paragraph (BBB)
        assert len(para_elems) == 1

        # Replace should proceed
        if len(para_elems) == 1:
            target_para = real_runs[0].get('para_elem')
            result = applier._apply_replace(
                target_para, violation_text, "Good text",
                "Test reason",
                runs_info, match_start,
                get_test_author(applier)
            )
            assert result == 'success'
    
    def test_delete_with_cross_para_content_fallbacks(self):
        """Match actually spans multiple paragraphs → fallback to comment"""
        applier = create_mock_applier()
        
        # Create paragraphs
        tc = create_table_cell_with_paragraphs(
            ['First line', 'Second line'],
            ['AAA', 'BBB']
        )
        
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        for p in tc.findall('.//w:p', NSMAP):
            body.append(p)
        applier.body_elem = body
        
        para_aaa = body.find(f'.//w:p[@w14:paraId="AAA"]', NSMAP)

        # Collect runs
        runs_info, combined_text, is_cross_paragraph, boundary_error = applier._collect_runs_info_across_paragraphs(para_aaa, 'BBB')

        assert boundary_error is None
        # Match spans both paragraphs (includes \n)
        violation_text = "First line\nSecond"
        match_start = 0
        match_end = match_start + len(violation_text)
        
        # Find affected runs
        affected = applier._find_affected_runs(runs_info, match_start, match_end)
        real_runs = [r for r in affected if not r.get('is_para_boundary', False)]
        
        # Check actual paragraphs spanned
        para_elems = set(r.get('para_elem') for r in real_runs if r.get('para_elem') is not None)
        
        # Should span 2 paragraphs
        assert len(para_elems) == 2
        
        # This would trigger fallback in actual code
        # We verify the detection logic works correctly


# ============================================================
# Tests: Table detection and JSON format
# ============================================================

def create_table_with_cells(cell_contents: list, row_para_ids: list = None) -> etree.Element:
    """
    Create a table element with one row and multiple cells.

    Args:
        cell_contents: List of lists - [[cell1_para1, cell1_para2], [cell2_para1], ...]
                       Each inner list contains paragraph texts for that cell
        row_para_ids: Flat list of paraIds for all paragraphs in order

    Returns:
        Table element (w:tbl)
    """
    tbl = etree.Element(f'{{{NS["w"]}}}tbl', nsmap=NSMAP)

    # Add tblGrid
    tbl_grid = etree.SubElement(tbl, f'{{{NS["w"]}}}tblGrid')
    for _ in cell_contents:
        etree.SubElement(tbl_grid, f'{{{NS["w"]}}}gridCol')

    # Create row
    tr = etree.SubElement(tbl, f'{{{NS["w"]}}}tr')

    para_idx = 0
    for cell_paras in cell_contents:
        tc = etree.SubElement(tr, f'{{{NS["w"]}}}tc')
        for para_text in cell_paras:
            p = etree.SubElement(tc, f'{{{NS["w"]}}}p')
            if row_para_ids and para_idx < len(row_para_ids):
                p.set(f'{{{NS["w14"]}}}paraId', row_para_ids[para_idx])
            r = etree.SubElement(p, f'{{{NS["w"]}}}r')
            t = etree.SubElement(r, f'{{{NS["w"]}}}t')
            t.text = para_text
            para_idx += 1

    return tbl


def create_multi_row_table(rows_data: list, para_ids: list = None) -> etree.Element:
    """
    Create a table element with multiple rows.

    Args:
        rows_data: List of rows, each row is a list of cell contents
                   [[[row1_cell1_para1], [row1_cell2_para1]], [[row2_cell1], [row2_cell2]]]
        para_ids: Flat list of paraIds for all paragraphs in document order

    Returns:
        Table element (w:tbl)
    """
    tbl = etree.Element(f'{{{NS["w"]}}}tbl', nsmap=NSMAP)

    # Add tblGrid (using first row's column count)
    if rows_data:
        tbl_grid = etree.SubElement(tbl, f'{{{NS["w"]}}}tblGrid')
        for _ in rows_data[0]:
            etree.SubElement(tbl_grid, f'{{{NS["w"]}}}gridCol')

    para_idx = 0
    for row_cells in rows_data:
        tr = etree.SubElement(tbl, f'{{{NS["w"]}}}tr')
        for cell_paras in row_cells:
            tc = etree.SubElement(tr, f'{{{NS["w"]}}}tc')
            for para_text in cell_paras:
                p = etree.SubElement(tc, f'{{{NS["w"]}}}p')
                if para_ids and para_idx < len(para_ids):
                    p.set(f'{{{NS["w14"]}}}paraId', para_ids[para_idx])
                r = etree.SubElement(p, f'{{{NS["w"]}}}r')
                t = etree.SubElement(r, f'{{{NS["w"]}}}t')
                t.text = para_text
                para_idx += 1

    return tbl


class TestTableDetection:
    """Tests for table detection helper methods"""

    def test_is_paragraph_in_table(self):
        """Paragraph inside table should be detected"""
        applier = create_mock_applier()

        # Create table with one cell
        tbl = create_table_with_cells([['Cell content']], ['AAA'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        # Find paragraph
        para = body.find('.//w:p', NSMAP)
        assert applier._is_paragraph_in_table(para) is True

    def test_is_paragraph_not_in_table(self):
        """Paragraph outside table should not be detected as in table"""
        applier = create_mock_applier()

        body = create_mock_body_with_paragraphs(['AAA'])
        applier.body_elem = body

        para = body.find('.//w:p', NSMAP)
        assert applier._is_paragraph_in_table(para) is False

    def test_find_ancestor_cell(self):
        """Should find the cell containing a paragraph"""
        applier = create_mock_applier()

        tbl = create_table_with_cells([['Cell1'], ['Cell2']], ['AAA', 'BBB'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)
        cell = applier._find_ancestor_cell(para)
        assert cell is not None
        assert cell.tag == f'{{{NS["w"]}}}tc'

    def test_find_ancestor_row(self):
        """Should find the row containing a paragraph"""
        applier = create_mock_applier()

        tbl = create_table_with_cells([['Cell1'], ['Cell2']], ['AAA', 'BBB'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)
        row = applier._find_ancestor_row(para)
        assert row is not None
        assert row.tag == f'{{{NS["w"]}}}tr'


class TestTableBoundaryDetection:
    """Tests for boundary detection in _collect_runs_info_across_paragraphs"""

    def test_body_to_table_boundary_crossed(self):
        """Crossing from body to table should return boundary_crossed error"""
        applier = create_mock_applier()

        # Create body with paragraph followed by table
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)

        # Add body paragraph
        p = etree.SubElement(body, f'{{{NS["w"]}}}p')
        p.set(f'{{{NS["w14"]}}}paraId', 'AAA')
        r = etree.SubElement(p, f'{{{NS["w"]}}}r')
        t = etree.SubElement(r, f'{{{NS["w"]}}}t')
        t.text = 'Body text'

        # Add table
        tbl = create_table_with_cells([['Table cell']], ['BBB'])
        body.append(tbl)

        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, 'BBB'
        )

        assert boundary_error == 'boundary_crossed'

    def test_table_to_body_boundary_crossed(self):
        """Crossing from table to body should return boundary_crossed error"""
        applier = create_mock_applier()

        # Create body with table followed by body paragraph
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)

        # Add table first
        tbl = create_table_with_cells([['Table cell']], ['AAA'])
        body.append(tbl)

        # Add body paragraph after table
        p = etree.SubElement(body, f'{{{NS["w"]}}}p')
        p.set(f'{{{NS["w14"]}}}paraId', 'BBB')
        r = etree.SubElement(p, f'{{{NS["w"]}}}r')
        t = etree.SubElement(r, f'{{{NS["w"]}}}t')
        t.text = 'Body text after table'

        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, 'BBB'
        )

        assert boundary_error == 'boundary_crossed'

    def test_table_row_boundary_collected(self):
        """Multi-row table content should be collected (row boundary checked later)"""
        applier = create_mock_applier()

        # Create table with two rows
        tbl = create_multi_row_table(
            [[['Row1 Cell1'], ['Row1 Cell2']], [['Row2 Cell1'], ['Row2 Cell2']]],
            ['AAA', 'BBB', 'CCC', 'DDD']
        )
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)

        # Content from row 1 to row 2 should be collected (no upfront error)
        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, 'CCC'
        )

        # No upfront boundary error - content is collected across rows
        assert boundary_error is None
        # Content from both rows should be present
        assert 'Row1 Cell1' in combined_text
        assert 'Row2 Cell1' in combined_text
        # Row boundary marker should be present
        assert '"], ["' in combined_text

        # _check_cross_row_boundary should detect if a match spans rows
        # Get runs that span from row 1 to row 2
        all_affected = applier._find_affected_runs(runs_info, 0, len(combined_text))
        real_runs = [r for r in all_affected
                     if not r.get('is_json_boundary', False)
                     and not r.get('is_json_escape', False)]
        assert applier._check_cross_row_boundary(real_runs) is True

    def test_same_row_no_boundary_error(self):
        """Same row content should not trigger boundary error"""
        applier = create_mock_applier()

        # Create table with one row, multiple cells
        tbl = create_table_with_cells([['Cell1'], ['Cell2'], ['Cell3']], ['AAA', 'BBB', 'CCC'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, 'CCC'
        )

        assert boundary_error is None


class TestTableJsonFormat:
    """Tests for JSON format in table row content collection"""

    def test_single_cell_json_format(self):
        """Single cell should produce JSON format ["content"]"""
        applier = create_mock_applier()

        tbl = create_table_with_cells([['Hello']], ['AAA'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, 'AAA'
        )

        assert boundary_error is None
        assert combined_text == '["Hello"]'

    def test_multi_cell_same_row_json_format(self):
        """Multiple cells in same row should produce JSON format ["cell1", "cell2"]"""
        applier = create_mock_applier()

        tbl = create_table_with_cells([['Cell1'], ['Cell2'], ['Cell3']], ['AAA', 'BBB', 'CCC'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, 'CCC'
        )

        assert boundary_error is None
        assert combined_text == '["Cell1", "Cell2", "Cell3"]'

    def test_multi_para_in_cell_json_format(self):
        """Multiple paragraphs in same cell should use \\n separator"""
        applier = create_mock_applier()

        # Cell with two paragraphs
        tbl = create_table_with_cells([['Para1', 'Para2']], ['AAA', 'BBB'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, 'BBB'
        )

        assert boundary_error is None
        # In JSON, newlines are escaped as \n
        assert combined_text == '["Para1\\nPara2"]'

    def test_json_escape_quotes(self):
        """Content with quotes should be JSON-escaped"""
        applier = create_mock_applier()

        tbl = create_table_with_cells([['He said "hello"']], ['AAA'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, 'AAA'
        )

        assert boundary_error is None
        # The quotes inside should be escaped
        assert combined_text == '["He said \\"hello\\""]'


class TestCrossCellBoundary:
    """Tests for cross-cell boundary detection"""

    def test_check_cross_cell_boundary_single_cell(self):
        """Runs from single cell should not trigger cross-cell"""
        applier = create_mock_applier()

        tbl = create_table_with_cells([['Cell content']], ['AAA'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        start_para = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)
        cell = applier._find_ancestor_cell(start_para)

        # Simulate run info with cell_elem
        runs_info = [
            {'text': 'Cell content', 'cell_elem': cell, 'start': 2, 'end': 14}
        ]

        assert applier._check_cross_cell_boundary(runs_info) is False

    def test_check_cross_cell_boundary_multiple_cells(self):
        """Runs from multiple cells should trigger cross-cell"""
        applier = create_mock_applier()

        tbl = create_table_with_cells([['Cell1'], ['Cell2']], ['AAA', 'BBB'])
        body = etree.Element(f'{{{NS["w"]}}}body', nsmap=NSMAP)
        body.append(tbl)
        applier.body_elem = body

        para_aaa = body.find('.//w:p[@w14:paraId="AAA"]', NSMAP)
        para_bbb = body.find('.//w:p[@w14:paraId="BBB"]', NSMAP)
        cell_aaa = applier._find_ancestor_cell(para_aaa)
        cell_bbb = applier._find_ancestor_cell(para_bbb)

        # Simulate run info spanning two cells
        runs_info = [
            {'text': 'Cell1', 'cell_elem': cell_aaa, 'start': 2, 'end': 7},
            {'text': 'Cell2', 'cell_elem': cell_bbb, 'start': 12, 'end': 17}
        ]

        assert applier._check_cross_cell_boundary(runs_info) is True


# ============================================================
# Tests: Real Document Table (tests/test.docx)
# ============================================================

class TestRealDocumentTable:
    """Tests using the real test.docx file for table operations"""

    @staticmethod
    def get_test_doc_path():
        """Get path to test.docx"""
        return Path(__file__).parent / 'test.docx'

    def test_table_detection_in_real_document(self):
        """Test that paragraphs in test.docx table are correctly detected"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Find a paragraph in the table (Row 1, Col 2: "全自动贴片生产线")
        # ParaId: 04B34894
        para = applier._find_para_node_by_id('04B34894')
        if para is None:
            pytest.skip("ParaId 04B34894 not found in document")

        assert applier._is_paragraph_in_table(para) is True

        # Check that we can find the cell and row
        cell = applier._find_ancestor_cell(para)
        assert cell is not None

        row = applier._find_ancestor_row(para)
        assert row is not None

        table = applier._find_ancestor_table(para)
        assert table is not None

    def test_same_row_cells_no_row_boundary_error(self):
        """Test that cells in the same row don't trigger row boundary error"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 1: 751B37B6 (序号), 30353808 (项目), 04B34894 (设备工装需求), 6CF3F5AF (数量), 662522F3 (备注)
        # Use start: 30353808, end: 6CF3F5AF (same row)
        start_para = applier._find_para_node_by_id('30353808')
        if start_para is None:
            pytest.skip("ParaId 30353808 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '6CF3F5AF'
        )

        # Should not have boundary error for same row
        assert boundary_error is None
        assert is_cross_para is True  # Table mode is always cross-paragraph

    def test_different_rows_collected_with_row_marker(self):
        """Test that content from different rows is collected with row boundary marker"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 1 cell: 04B34894 (全自动贴片生产线)
        # Row 2 cell: 2B402507 (恒温干燥柜)
        start_para = applier._find_para_node_by_id('04B34894')
        if start_para is None:
            pytest.skip("ParaId 04B34894 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '2B402507'
        )

        # No upfront boundary error - content is collected across rows
        assert boundary_error is None

        # Content from both rows should be present
        assert '全自动贴片生产线' in combined_text
        assert '恒温干燥柜' in combined_text

        # Row boundary marker should be present
        assert '"], ["' in combined_text

        # _check_cross_row_boundary should detect the row span
        all_affected = applier._find_affected_runs(runs_info, 0, len(combined_text))
        real_runs = [r for r in all_affected
                     if not r.get('is_json_boundary', False)
                     and not r.get('is_json_escape', False)]
        assert applier._check_cross_row_boundary(real_runs) is True

    def test_json_format_for_real_table_row(self):
        """Test JSON format generation for a real table row"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 0 (header): 12C551B6 -> 5CC914F0
        # Expected: ["序号", "项目", "设备工装需求", "数量", "备注"]
        start_para = applier._find_para_node_by_id('12C551B6')
        if start_para is None:
            pytest.skip("ParaId 12C551B6 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '5CC914F0'
        )

        assert boundary_error is None
        # Check JSON format
        assert combined_text.startswith('["')
        assert combined_text.endswith('"]')
        assert '", "' in combined_text  # Cell separator

        # Verify content is present
        assert '序号' in combined_text
        assert '项目' in combined_text
        assert '设备工装需求' in combined_text
        assert '数量' in combined_text
        assert '备注' in combined_text

    def test_single_cell_in_real_table(self):
        """Test collecting runs from a single cell in real table"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Single cell: Row 1, Col 2 - "全自动贴片生产线" (paraId: 04B34894)
        start_para = applier._find_para_node_by_id('04B34894')
        if start_para is None:
            pytest.skip("ParaId 04B34894 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '04B34894'  # Same start and end
        )

        assert boundary_error is None
        # Single cell should produce ["content"]
        assert combined_text == '["全自动贴片生产线"]'

    def test_cross_cell_detection_in_real_table(self):
        """Test cross-cell boundary detection with real table"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 1: cells with paraIds 30353808, 04B34894, 6CF3F5AF
        # Collect runs spanning these cells
        start_para = applier._find_para_node_by_id('30353808')
        if start_para is None:
            pytest.skip("ParaId 30353808 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '6CF3F5AF'
        )

        assert boundary_error is None

        # Filter to get real runs (not JSON boundaries)
        real_runs = [r for r in runs_info
                     if not r.get('is_json_boundary', False)
                     and not r.get('is_json_escape', False)]

        # Should span multiple cells
        assert applier._check_cross_cell_boundary(real_runs) is True


class TestRealDocumentTableViolation:
    """Tests for processing violations in real document tables"""

    @staticmethod
    def get_test_doc_path():
        """Get path to test.docx"""
        return Path(__file__).parent / 'test.docx'

    def test_single_cell_violation_match(self):
        """Test matching violation text within a single cell"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 2, Col 2: "恒温干燥柜" (paraId: 2B402507)
        start_para = applier._find_para_node_by_id('2B402507')
        if start_para is None:
            pytest.skip("ParaId 2B402507 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '2B402507'
        )

        assert boundary_error is None

        # Search for "干燥" within the cell
        violation_text = '干燥'
        pos = combined_text.find(violation_text)

        # Should find it (accounting for JSON format ["恒温干燥柜"])
        assert pos != -1, f"'{violation_text}' not found in '{combined_text}'"

    def test_cross_cell_violation_in_same_row(self):
        """Test that violation spanning multiple cells in same row can be found"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 0 (header): all cells
        start_para = applier._find_para_node_by_id('12C551B6')
        if start_para is None:
            pytest.skip("ParaId 12C551B6 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '5CC914F0'
        )

        assert boundary_error is None

        # A cross-cell violation text would include the cell separator
        # For example: '项目", "设备工装需求'
        cross_cell_text = '项目", "设备工装需求'
        pos = combined_text.find(cross_cell_text)

        assert pos != -1, f"Cross-cell text '{cross_cell_text}' not found in '{combined_text}'"

        # Find affected runs
        match_end = pos + len(cross_cell_text)
        affected = applier._find_affected_runs(runs_info, pos, match_end)

        # Filter real runs
        real_runs = [r for r in affected
                     if not r.get('is_json_boundary', False)
                     and not r.get('is_json_escape', False)]

        # Should detect cross-cell boundary
        assert applier._check_cross_cell_boundary(real_runs) is True


# ============================================================
# Test Class: Filter Real Runs (Issue 1 - JSON boundary runs without elem/rPr)
# ============================================================

class TestFilterRealRuns:
    """Test that _filter_real_runs correctly filters out synthetic boundary markers"""

    def test_filter_removes_json_boundary_runs(self):
        """Test that JSON boundary runs (["  and "]) are filtered out"""
        applier = create_mock_applier()

        runs = [
            {'text': '["', 'start': 0, 'end': 2, 'is_json_boundary': True},
            {'text': 'hello', 'start': 2, 'end': 7, 'elem': 'mock_elem', 'rPr': None},
            {'text': '"]', 'start': 7, 'end': 9, 'is_json_boundary': True},
        ]

        filtered = applier._filter_real_runs(runs)

        assert len(filtered) == 1
        assert filtered[0]['text'] == 'hello'
        assert 'elem' in filtered[0]

    def test_filter_removes_cell_boundary_runs(self):
        """Test that cell boundary runs (", ") are filtered out"""
        applier = create_mock_applier()

        runs = [
            {'text': '["', 'start': 0, 'end': 2, 'is_json_boundary': True},
            {'text': 'cell1', 'start': 2, 'end': 7, 'elem': 'mock1', 'rPr': None},
            {'text': '", "', 'start': 7, 'end': 11, 'is_json_boundary': True, 'is_cell_boundary': True},
            {'text': 'cell2', 'start': 11, 'end': 16, 'elem': 'mock2', 'rPr': None},
            {'text': '"]', 'start': 16, 'end': 18, 'is_json_boundary': True},
        ]

        filtered = applier._filter_real_runs(runs)

        assert len(filtered) == 2
        assert filtered[0]['text'] == 'cell1'
        assert filtered[1]['text'] == 'cell2'

    def test_filter_removes_para_boundary_runs(self):
        """Test that paragraph boundary runs are filtered out"""
        applier = create_mock_applier()

        runs = [
            {'text': 'para1', 'start': 0, 'end': 5, 'elem': 'mock1', 'rPr': None},
            {'text': '\n', 'start': 5, 'end': 6, 'is_para_boundary': True},
            {'text': 'para2', 'start': 6, 'end': 11, 'elem': 'mock2', 'rPr': None},
        ]

        filtered = applier._filter_real_runs(runs)

        assert len(filtered) == 2
        assert filtered[0]['text'] == 'para1'
        assert filtered[1]['text'] == 'para2'

    def test_filter_removes_json_escape_runs(self):
        """Test that JSON escape marker runs are filtered out"""
        applier = create_mock_applier()

        runs = [
            {'text': 'before', 'start': 0, 'end': 6, 'elem': 'mock1', 'rPr': None},
            {'text': '\\', 'start': 6, 'end': 7, 'is_json_escape': True},
            {'text': 'after', 'start': 7, 'end': 12, 'elem': 'mock2', 'rPr': None},
        ]

        filtered = applier._filter_real_runs(runs)

        assert len(filtered) == 2
        assert filtered[0]['text'] == 'before'
        assert filtered[1]['text'] == 'after'

    def test_filter_keeps_all_real_runs(self):
        """Test that real runs with elem are preserved"""
        applier = create_mock_applier()

        runs = [
            {'text': 'run1', 'start': 0, 'end': 4, 'elem': 'mock1', 'rPr': None},
            {'text': 'run2', 'start': 4, 'end': 8, 'elem': 'mock2', 'rPr': 'mock_rPr'},
            {'text': 'run3', 'start': 8, 'end': 12, 'elem': 'mock3', 'rPr': None, 'is_drawing': True},
        ]

        filtered = applier._filter_real_runs(runs)

        assert len(filtered) == 3
        assert all('elem' in r for r in filtered)


# ============================================================
# Test Class: Original Text Helpers (Issue 2 - JSON escaped text for mutations)
# ============================================================

class TestOriginalTextHelpers:
    """Test helper methods for handling JSON-escaped vs original text"""

    def test_get_run_original_text_with_original(self):
        """Test that _get_run_original_text returns original_text when present"""
        applier = create_mock_applier()

        run = {
            'text': 'He said \\"hello\\"',  # JSON escaped
            'original_text': 'He said "hello"',  # Original
            'start': 0,
            'end': 17
        }

        result = applier._get_run_original_text(run)
        assert result == 'He said "hello"'

    def test_get_run_original_text_without_original(self):
        """Test that _get_run_original_text returns text when no original_text"""
        applier = create_mock_applier()

        run = {
            'text': 'simple text',
            'start': 0,
            'end': 11
        }

        result = applier._get_run_original_text(run)
        assert result == 'simple text'

    def test_translate_escaped_offset_no_escaping(self):
        """Test offset translation when no escaping present"""
        applier = create_mock_applier()

        run = {
            'text': 'hello world',
            'start': 0,
            'end': 11
        }

        # No original_text means no escaping, offset unchanged
        assert applier._translate_escaped_offset(run, 0) == 0
        assert applier._translate_escaped_offset(run, 5) == 5
        assert applier._translate_escaped_offset(run, 11) == 11

    def test_translate_escaped_offset_with_quote_escaping(self):
        """Test offset translation with escaped quotes"""
        applier = create_mock_applier()

        # Original: He said "hi"
        # Escaped:  He said \"hi\"
        run = {
            'text': 'He said \\"hi\\"',  # 14 chars
            'original_text': 'He said "hi"',  # 12 chars
            'start': 0,
            'end': 14
        }

        # Before the first escape
        assert applier._translate_escaped_offset(run, 0) == 0
        assert applier._translate_escaped_offset(run, 7) == 7  # 'He said'

        # After both escapes (end of string)
        assert applier._translate_escaped_offset(run, 14) == 12

    def test_translate_escaped_offset_with_backslash_escaping(self):
        """Test offset translation with escaped backslashes"""
        applier = create_mock_applier()

        # Original: path\to\file
        # Escaped:  path\\to\\file
        run = {
            'text': 'path\\\\to\\\\file',  # 14 chars
            'original_text': 'path\\to\\file',  # 12 chars
            'start': 0,
            'end': 14
        }

        assert applier._translate_escaped_offset(run, 0) == 0
        assert applier._translate_escaped_offset(run, 4) == 4  # 'path'
        assert applier._translate_escaped_offset(run, 14) == 12

    def test_translate_escaped_offset_with_newline_escaping(self):
        """Test offset translation with escaped newlines"""
        applier = create_mock_applier()

        # Original: line1\nline2 (11 chars with actual newline)
        # Escaped:  line1\\nline2 (12 chars with \n as two chars)
        run = {
            'text': 'line1\\nline2',  # 12 chars
            'original_text': 'line1\nline2',  # 11 chars
            'start': 0,
            'end': 12
        }

        assert applier._translate_escaped_offset(run, 0) == 0
        assert applier._translate_escaped_offset(run, 5) == 5  # 'line1'
        assert applier._translate_escaped_offset(run, 12) == 11

    def test_translate_escaped_offset_boundary_cases(self):
        """Test boundary cases for offset translation"""
        applier = create_mock_applier()

        run = {
            'text': 'test\\"text',  # 10 chars
            'original_text': 'test"text',  # 9 chars
            'start': 0,
            'end': 10
        }

        # Negative offset returns 0
        assert applier._translate_escaped_offset(run, -1) == 0

        # Offset beyond end returns original length
        assert applier._translate_escaped_offset(run, 100) == 9

    def test_is_table_mode_with_original_text(self):
        """Test _is_table_mode returns True when runs have original_text"""
        applier = create_mock_applier()

        runs = [
            {'text': 'escaped', 'original_text': 'original', 'start': 0, 'end': 7},
            {'text': 'normal', 'start': 7, 'end': 13}
        ]

        assert applier._is_table_mode(runs) is True

    def test_is_table_mode_without_original_text(self):
        """Test _is_table_mode returns False when no runs have original_text"""
        applier = create_mock_applier()

        runs = [
            {'text': 'normal1', 'start': 0, 'end': 7},
            {'text': 'normal2', 'start': 7, 'end': 14}
        ]

        assert applier._is_table_mode(runs) is False

    def test_decode_json_escaped_quotes(self):
        """Test _decode_json_escaped decodes escaped quotes"""
        applier = create_mock_applier()

        assert applier._decode_json_escaped('He said \\"hello\\"') == 'He said "hello"'

    def test_decode_json_escaped_backslash(self):
        """Test _decode_json_escaped decodes escaped backslashes"""
        applier = create_mock_applier()

        assert applier._decode_json_escaped('path\\\\to\\\\file') == 'path\\to\\file'

    def test_decode_json_escaped_newline(self):
        """Test _decode_json_escaped decodes escaped newlines"""
        applier = create_mock_applier()

        assert applier._decode_json_escaped('line1\\nline2') == 'line1\nline2'

    def test_decode_json_escaped_empty(self):
        """Test _decode_json_escaped handles empty string"""
        applier = create_mock_applier()

        assert applier._decode_json_escaped('') == ''

    def test_decode_json_escaped_no_escapes(self):
        """Test _decode_json_escaped returns unchanged text with no escapes"""
        applier = create_mock_applier()

        assert applier._decode_json_escaped('normal text') == 'normal text'


# ============================================================
# Test Class: Apply Methods with Escaped Text (Issue 2)
# ============================================================

class TestApplyMethodsWithEscapedText:
    """Test that apply methods correctly use original text for mutations"""

    def test_apply_delete_with_escaped_quote(self):
        """Test _apply_delete correctly decodes JSON-escaped violation_text"""
        # Create paragraph with text containing quotes
        p = create_paragraph_xml('He said "hello" to me')
        applier = create_mock_applier()

        # Simulate table mode where run has both escaped and original text
        runs_info = [{
            'text': 'He said \\"hello\\" to me',  # JSON escaped (23 chars)
            'original_text': 'He said "hello" to me',  # Original (21 chars)
            'start': 0,
            'end': 23,
            'elem': p.find('.//w:r', namespaces={'w': NS['w']}),
            'rPr': None,
            'para_elem': p
        }]

        # In table mode, LLM sees JSON-escaped content, so it returns escaped text
        # violation_text is what the LLM returns (escaped form)
        violation_text = '\\"hello\\"'  # LLM returns escaped version

        # Find position in the escaped combined text
        combined_text = runs_info[0]['text']
        match_start = combined_text.find('\\"hello\\"')
        assert match_start != -1

        result = applier._apply_delete(
            p, violation_text, 'test reason',
            runs_info, match_start, 'test-author'
        )

        # Should succeed without corrupting the document
        assert result == 'success'

        # Verify the delText contains the DECODED text (actual quotes, not escaped)
        del_text_elem = p.find('.//w:delText', namespaces={'w': NS['w']})
        assert del_text_elem is not None
        # The delText should contain decoded text: "hello" not \"hello\"
        assert del_text_elem.text == '"hello"', f"delText should be decoded, got: {del_text_elem.text}"
        assert '\\' not in del_text_elem.text, "delText should not contain backslashes"

    def test_apply_replace_with_escaped_quote(self):
        """Test _apply_replace correctly decodes JSON-escaped text"""
        p = create_paragraph_xml('Value is "old" here')
        applier = create_mock_applier()

        runs_info = [{
            'text': 'Value is \\"old\\" here',  # JSON escaped (21 chars)
            'original_text': 'Value is "old" here',  # Original (19 chars)
            'start': 0,
            'end': 21,
            'elem': p.find('.//w:r', namespaces={'w': NS['w']}),
            'rPr': None,
            'para_elem': p
        }]

        # LLM returns escaped versions - replace entire "old" with "newval"
        # This tests that quotes at different positions get decoded properly
        violation_text = '\\"old\\"'
        revised_text = '\\"newval\\"'

        combined_text = runs_info[0]['text']
        match_start = combined_text.find('\\"old\\"')
        assert match_start != -1

        result = applier._apply_replace(
            p, violation_text, revised_text, 'test reason',
            runs_info, match_start, 'test-author'
        )

        assert result == 'success'

        # The diff between \"old\" and \"newval\" will have:
        # - equal: \" (decoded to ")
        # - delete: old
        # - insert: newval
        # - equal: \" (decoded to ")
        # So we check that the output doesn't contain backslash escapes

        # Get all text content from the paragraph
        all_text = etree.tostring(p, encoding='unicode')

        # Should not contain literal \\" sequences in the output
        assert '\\"' not in all_text or '&quot;' in all_text, \
            f"Output should not contain escaped quotes: {all_text[:200]}"

    def test_apply_replace_decodes_full_delete(self):
        """Test _apply_replace decodes when entire text is replaced"""
        p = create_paragraph_xml('Say "hi"')
        applier = create_mock_applier()

        runs_info = [{
            'text': 'Say \\"hi\\"',  # JSON escaped
            'original_text': 'Say "hi"',  # Original
            'start': 0,
            'end': 10,
            'elem': p.find('.//w:r', namespaces={'w': NS['w']}),
            'rPr': None,
            'para_elem': p
        }]

        # Replace entire "hi" with "bye" - different lengths force full delete/insert
        violation_text = '\\"hi\\"'
        revised_text = '\\"bye\\"'

        combined_text = runs_info[0]['text']
        match_start = combined_text.find('\\"hi\\"')
        assert match_start != -1

        result = applier._apply_replace(
            p, violation_text, revised_text, 'test reason',
            runs_info, match_start, 'test-author'
        )

        assert result == 'success'

        # Verify ins/del content doesn't have escaped backslashes
        ins_elems = p.findall('.//w:ins//w:t', namespaces={'w': NS['w']})
        for ins_elem in ins_elems:
            if ins_elem.text:
                assert '\\' not in ins_elem.text, \
                    f"Inserted text should not have backslashes: {ins_elem.text}"

    def test_apply_manual_with_escaped_quote(self):
        """Test _apply_manual correctly splits runs with quoted content"""
        p = create_paragraph_xml('Text "quoted" more')
        applier = create_mock_applier()

        runs_info = [{
            'text': 'Text \\"quoted\\" more',  # JSON escaped (20 chars)
            'original_text': 'Text "quoted" more',  # Original (18 chars)
            'start': 0,
            'end': 20,
            'elem': p.find('.//w:r', namespaces={'w': NS['w']}),
            'rPr': None,
            'para_elem': p
        }]

        violation_text = '"quoted"'

        combined_text = runs_info[0]['text']
        match_start = combined_text.find('\\"quoted\\"')
        assert match_start != -1

        result = applier._apply_manual(
            p, violation_text, 'test reason', 'suggestion',
            runs_info, match_start, 'test-author'
        )

        assert result == 'success'

        # Verify comment markers were inserted
        comment_start = p.find('.//w:commentRangeStart', namespaces={'w': NS['w']})
        comment_end = p.find('.//w:commentRangeEnd', namespaces={'w': NS['w']})
        assert comment_start is not None
        assert comment_end is not None

    def test_before_after_text_uses_original(self):
        """Test that before/after text splits use original (unescaped) text"""
        p = create_paragraph_xml('prefix "target" suffix')
        applier = create_mock_applier()

        # The run element
        run_elem = p.find('.//w:r', namespaces={'w': NS['w']})

        runs_info = [{
            'text': 'prefix \\"target\\" suffix',  # 24 chars escaped
            'original_text': 'prefix "target" suffix',  # 22 chars original
            'start': 0,
            'end': 24,
            'elem': run_elem,
            'rPr': None,
            'para_elem': p
        }]

        # Delete "target" (with quotes)
        violation_text = '"target"'

        # Position in escaped text
        combined_text = runs_info[0]['text']
        match_start = combined_text.find('\\"target\\"')

        result = applier._apply_delete(
            p, violation_text, 'test reason',
            runs_info, match_start, 'test-author'
        )

        assert result == 'success'

        # Get all text runs to verify splits
        all_runs = p.findall('.//w:r', namespaces={'w': NS['w']})

        # Find the before text (should be 'prefix ' not 'prefix \\')
        texts = []
        for run in all_runs:
            t = run.find('w:t', namespaces={'w': NS['w']})
            if t is not None and t.text:
                texts.append(t.text)

        # Should have correct unescaped prefix and suffix
        # 'prefix ' and ' suffix' should be present as separate runs
        all_text = ''.join(texts)
        assert 'prefix ' in all_text or all_text.startswith('prefix')
        assert ' suffix' in all_text or all_text.endswith('suffix')

        # Should NOT have escaped backslashes in the output
        assert '\\' not in all_text, f"Output should not contain backslashes: {all_text}"


# ============================================================
# Test Class: Apply Methods with JSON Boundary Runs (Issue 1)
# ============================================================

class TestApplyMethodsWithBoundaryRuns:
    """Test that apply methods don't fail when runs contain JSON boundary markers"""

    def test_apply_delete_filters_boundary_runs(self):
        """Test _apply_delete filters out JSON boundary runs before processing"""
        p = create_paragraph_xml('cell content')
        applier = create_mock_applier()

        run_elem = p.find('.//w:r', namespaces={'w': NS['w']})

        # Simulated table mode runs with JSON boundaries
        runs_info = [
            {'text': '["', 'start': 0, 'end': 2, 'is_json_boundary': True, 'para_elem': p},
            {'text': 'cell content', 'start': 2, 'end': 14, 'elem': run_elem, 'rPr': None, 'para_elem': p},
            {'text': '"]', 'start': 14, 'end': 16, 'is_json_boundary': True, 'para_elem': p},
        ]

        # Delete "cell" - note we need to account for the JSON prefix
        violation_text = 'cell'
        match_start = 2  # After '["'

        result = applier._apply_delete(
            p, violation_text, 'test reason',
            runs_info, match_start, 'test-author'
        )

        # Should succeed without KeyError on boundary runs
        assert result == 'success'

    def test_apply_replace_filters_boundary_runs(self):
        """Test _apply_replace filters out JSON boundary runs before processing"""
        p = create_paragraph_xml('old value')
        applier = create_mock_applier()

        run_elem = p.find('.//w:r', namespaces={'w': NS['w']})

        runs_info = [
            {'text': '["', 'start': 0, 'end': 2, 'is_json_boundary': True, 'para_elem': p},
            {'text': 'old value', 'start': 2, 'end': 11, 'elem': run_elem, 'rPr': None, 'para_elem': p},
            {'text': '"]', 'start': 11, 'end': 13, 'is_json_boundary': True, 'para_elem': p},
        ]

        violation_text = 'old'
        revised_text = 'new'
        match_start = 2

        result = applier._apply_replace(
            p, violation_text, revised_text, 'test reason',
            runs_info, match_start, 'test-author'
        )

        assert result == 'success'

    def test_apply_manual_filters_boundary_runs(self):
        """Test _apply_manual filters out JSON boundary runs before processing"""
        p = create_paragraph_xml('target text')
        applier = create_mock_applier()

        run_elem = p.find('.//w:r', namespaces={'w': NS['w']})

        runs_info = [
            {'text': '["', 'start': 0, 'end': 2, 'is_json_boundary': True, 'para_elem': p},
            {'text': 'target text', 'start': 2, 'end': 13, 'elem': run_elem, 'rPr': None, 'para_elem': p},
            {'text': '"]', 'start': 13, 'end': 15, 'is_json_boundary': True, 'para_elem': p},
        ]

        violation_text = 'target'
        match_start = 2

        result = applier._apply_manual(
            p, violation_text, 'test reason', 'suggestion',
            runs_info, match_start, 'test-author'
        )

        assert result == 'success'

    def test_apply_delete_with_only_boundary_runs_returns_fallback(self):
        """Test that _apply_delete returns fallback if all runs are boundaries"""
        p = create_paragraph_xml('x')
        applier = create_mock_applier()

        # All runs are boundary markers (no real runs)
        runs_info = [
            {'text': '["', 'start': 0, 'end': 2, 'is_json_boundary': True, 'para_elem': p},
            {'text': '", "', 'start': 2, 'end': 6, 'is_json_boundary': True, 'is_cell_boundary': True, 'para_elem': p},
            {'text': '"]', 'start': 6, 'end': 8, 'is_json_boundary': True, 'para_elem': p},
        ]

        # Try to delete something that spans only boundaries
        violation_text = '["'
        match_start = 0

        result = applier._apply_delete(
            p, violation_text, 'test reason',
            runs_info, match_start, 'test-author'
        )

        # Should fallback since no real runs to modify
        assert result == 'fallback'


# ============================================================
# Main
# ============================================================

if __name__ == '__main__':
    pytest.main([__file__, '-v'])
