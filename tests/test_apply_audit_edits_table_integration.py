#!/usr/bin/env python3
"""
Integration-style tests for table operations in apply_audit_edits.py.
"""

import sys
import tempfile
from pathlib import Path
import json

TESTS_DIR = Path(__file__).parent
sys.path.insert(0, str(TESTS_DIR))

from lxml import etree
from docx import Document

from _apply_audit_edits_helpers import AuditEditApplier, NS


def _add_para_ids(docx_path: Path):
    """Add w14:paraId attributes to all paragraphs (simulate Word 2013+ behavior)."""
    doc = Document(str(docx_path))
    body_elem = doc._element.body

    # Add w14 namespace if not present
    nsmap = body_elem.nsmap
    if 'w14' not in nsmap:
        # Register namespace at document level
        etree.register_namespace('w14', NS['w14'])

    counter = 0
    for p in body_elem.iter(f'{{{NS["w"]}}}p'):
        # Generate unique 8-character hex ID (simulating Word's format)
        para_id = f'{counter:08X}'
        p.set('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', para_id)
        counter += 1

    # Save with updated paraId attributes
    doc.save(str(docx_path))


def _get_para_ids(docx_path: Path) -> dict:
    """Extract paragraph IDs from the document for testing."""
    doc = Document(str(docx_path))
    body_elem = doc._element.body

    para_ids = {}
    for p in body_elem.iter(f'{{{NS["w"]}}}p'):
        para_id = p.get('{http://schemas.microsoft.com/office/word/2010/wordml}paraId')
        if para_id:
            # Extract text content
            text = ''.join(t.text or '' for t in p.iter(f'{{{NS["w"]}}}t'))
            para_ids[text] = para_id

    return para_ids


def create_multi_cell_table_document(docx_path: Path):
    """
    Create a test document with a table for multi-cell operations.

    Table structure:
    +-------+-------+-------+
    | Row 1 | Data1 | Data2 |
    +-------+-------+-------+
    | Row 2 | Data3 | Data4 |
    +-------+-------+-------+
    | Row 3 | Data5 | Data6 |
    +-------+-------+-------+
    """
    doc = Document()

    # Add title
    doc.add_heading('Multi-Cell Operations Test', 0)

    # Create table with 3 rows x 3 columns
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # Fill table data
    cells = [
        ['Row 1', 'Data1', 'Data2'],
        ['Row 2', 'Data3', 'Data4'],
        ['Row 3', 'Data5', 'Data6']
    ]

    for i, row_data in enumerate(cells):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text

    # Save document
    doc.save(str(docx_path))

    # Add paraId attributes (required for apply_audit_edits.py)
    _add_para_ids(docx_path)

    return doc


def create_boundary_marker_table_document(docx_path: Path):
    """
    Create a test document with a table for boundary marker operations.

    Table structure:
    +-------+-------+-------+
    | Cell1 | Cell2 | Cell3 |
    +-------+-------+-------+
    | Row2A | Row2B | Row2C |
    +-------+-------+-------+
    """
    doc = Document()

    # Add title
    doc.add_heading('Boundary Marker Operations Test', 0)

    # Create table with 2 rows x 3 columns
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    # Fill table data
    cells = [
        ['Cell1', 'Cell2', 'Cell3'],
        ['Row2A', 'Row2B', 'Row2C']
    ]

    for i, row_data in enumerate(cells):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text

    # Save document
    doc.save(str(docx_path))

    # Add paraId attributes (required for apply_audit_edits.py)
    _add_para_ids(docx_path)

    return doc


def test_multi_row_delete(tmp_path):
    """Test deleting content across multiple table rows."""
    print("\n" + "=" * 60)
    print("TEST: Multi-row Delete")
    print("=" * 60)

    # Create test document
    docx_path = tmp_path / 'test_multi_row_delete.docx'
    create_multi_cell_table_document(docx_path)

    # Get paragraph IDs
    para_ids = _get_para_ids(docx_path)

    # Create JSONL with multi-row delete (JSON format)
    # Delete first two rows: ["Row 1", "Data1", "Data2"], ["Row 2", "Data3", "Data4"]
    jsonl_path = tmp_path / 'test_multi_row_delete.jsonl'

    # Find start and end para IDs
    uuid_start = para_ids.get('Row 1')
    uuid_end = para_ids.get('Data4')

    if not uuid_start or not uuid_end:
        print("ERROR: Could not find paragraph IDs")
        print(f"Available IDs: {para_ids}")
        assert False, "Could not find required paragraph IDs"

    with open(jsonl_path, 'w', encoding='utf-8') as f:
        # Meta line (required)
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')

        # Edit item: delete multi-row content
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '["Row 1", "Data1", "Data2"], ["Row 2", "Data3", "Data4"]',
            'violation_reason': 'Multi-row delete test',
            'fix_action': 'delete',
            'revised_text': '',
            'category': 'test',
            'rule_id': 'TEST001',
            'heading': 'Test'
        }
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')

    # Apply edits
    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()

        # Check results
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, (
            f"Multi-row delete failed: {results[0].error_message if results else 'No results'}"
        )

        print("\n✓ Multi-row delete succeeded")
        print(f"  Output: {applier.output_path}")

    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_multi_cell_same_row_delete(tmp_path):
    """Test deleting content across multiple cells in the same row."""
    print("\n" + "=" * 60)
    print("TEST: Multi-cell Same Row Delete")
    print("=" * 60)

    # Create test document
    docx_path = tmp_path / 'test_multi_cell_delete.docx'
    create_multi_cell_table_document(docx_path)

    # Get paragraph IDs
    para_ids = _get_para_ids(docx_path)

    # Delete cells in first row: ["Data1", "Data2"] (skipping "Row 1")
    jsonl_path = tmp_path / 'test_multi_cell_delete.jsonl'

    uuid_start = para_ids.get('Data1')
    uuid_end = para_ids.get('Data2')

    if not uuid_start or not uuid_end:
        print("ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"

    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')

        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '"Data1", "Data2"',  # Same row, two cells
            'violation_reason': 'Multi-cell same row delete test',
            'fix_action': 'delete',
            'revised_text': '',
            'category': 'test',
            'rule_id': 'TEST002',
            'heading': 'Test'
        }
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')

    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()

        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, (
            f"Multi-cell delete failed: {results[0].error_message if results else 'No results'}"
        )

        print("\n✓ Multi-cell same row delete succeeded")
        print(f"  Output: {applier.output_path}")

    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_multi_row_replace(tmp_path):
    """Test replacing content across multiple table rows."""
    print("\n" + "=" * 60)
    print("TEST: Multi-row Replace")
    print("=" * 60)

    # Create test document
    docx_path = tmp_path / 'test_multi_row_replace.docx'
    create_multi_cell_table_document(docx_path)

    # Get paragraph IDs
    para_ids = _get_para_ids(docx_path)

    jsonl_path = tmp_path / 'test_multi_row_replace.jsonl'

    uuid_start = para_ids.get('Row 1')
    uuid_end = para_ids.get('Data4')

    if not uuid_start or not uuid_end:
        print("ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"

    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')

        # Replace first two rows with new content
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '["Row 1", "Data1", "Data2"], ["Row 2", "Data3", "Data4"]',
            'violation_reason': 'Multi-row replace test',
            'fix_action': 'replace',
            'revised_text': '["Row 1", "New1", "New2"], ["Row 2", "New3", "New4"]',
            'category': 'test',
            'rule_id': 'TEST003',
            'heading': 'Test'
        }
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')

    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()

        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, (
            f"Multi-row replace failed: {results[0].error_message if results else 'No results'}"
        )

        print("\n✓ Multi-row replace succeeded")
        print(f"  Output: {applier.output_path}")

    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_multi_cell_same_row_replace(tmp_path):
    """Test replacing content across multiple cells in the same row."""
    print("\n" + "=" * 60)
    print("TEST: Multi-cell Same Row Replace")
    print("=" * 60)

    # Create test document
    docx_path = tmp_path / 'test_multi_cell_replace.docx'
    create_multi_cell_table_document(docx_path)

    # Get paragraph IDs
    para_ids = _get_para_ids(docx_path)

    jsonl_path = tmp_path / 'test_multi_cell_replace.jsonl'

    uuid_start = para_ids.get('Data1')
    uuid_end = para_ids.get('Data2')

    if not uuid_start or not uuid_end:
        print("ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"

    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')

        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '"Data1", "Data2"',  # Same row, two cells
            'violation_reason': 'Multi-cell same row replace test',
            'fix_action': 'replace',
            'revised_text': '"New1", "New2"',
            'category': 'test',
            'rule_id': 'TEST004',
            'heading': 'Test'
        }
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')

    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()

        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, (
            f"Multi-cell replace failed: {results[0].error_message if results else 'No results'}"
        )

        print("\n✓ Multi-cell same row replace succeeded")
        print(f"  Output: {applier.output_path}")

    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_single_cell_replace_to_empty(tmp_path):
    """Test replacing a single cell with empty content via cross-cell match."""
    print("\n" + "=" * 60)
    print("TEST: Single Cell Replace to Empty")
    print("=" * 60)

    docx_path = tmp_path / 'test_single_cell_replace_empty.docx'
    create_multi_cell_table_document(docx_path)

    para_ids = _get_para_ids(docx_path)

    jsonl_path = tmp_path / 'test_single_cell_replace_empty.jsonl'

    uuid_start = para_ids.get('Data1')
    uuid_end = para_ids.get('Data2')

    if not uuid_start or not uuid_end:
        print("ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"

    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')

        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '"Data1", "Data2"',
            'violation_reason': 'Single-cell delete via replace test',
            'fix_action': 'replace',
            'revised_text': '"", "Data2"',
            'category': 'test',
            'rule_id': 'TEST005',
            'heading': 'Test'
        }
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')

    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()

        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, (
            f"Single-cell replace failed: {results[0].error_message if results else 'No results'}"
        )

        edited_doc = Document(str(applier.output_path))
        table = edited_doc.tables[0]
        assert table.rows[0].cells[1].text == "", "Expected Data1 cell to be empty"
        assert table.rows[0].cells[2].text == "Data2", "Expected Data2 cell to remain"

        print("\n✓ Single-cell replace to empty succeeded")
        print(f"  Output: {applier.output_path}")

    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_cell_separator_delete_fallback(tmp_path):
    """
    Test that attempting to delete a cell separator (\", \") falls back to comment.

    Scenario: Replace '[\"Cell1\", \"Cell2\"]' with '[\"Cell1Cell2\"]' (merge two cells)
    Expected: Fallback to comment because deleting \", \" lands on boundary marker
    """
    print("\n" + "=" * 60)
    print("TEST: Cell Separator Delete Fallback")
    print("=" * 60)

    # Create test document
    docx_path = tmp_path / 'test_cell_separator_delete.docx'
    create_boundary_marker_table_document(docx_path)

    # Get paragraph IDs
    para_ids = _get_para_ids(docx_path)

    # Create JSONL that would merge two cells (delete separator)
    jsonl_path = tmp_path / 'test_cell_separator_delete.jsonl'

    uuid_start = para_ids.get('Cell1')
    uuid_end = para_ids.get('Cell2')

    if not uuid_start or not uuid_end:
        print("ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"

    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')

        # Replace operation that would delete the ", " separator
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '"Cell1", "Cell2"',
            'violation_reason': 'Test cell merge - should fallback to comment',
            'fix_action': 'replace',
            'revised_text': '"Cell1Cell2"',  # Merge - removes ", " separator
            'category': 'test',
            'rule_id': 'TEST_BOUNDARY_001',
            'heading': 'Test'
        }
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')

    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()

        # Check results - should succeed with warning (fallback to comment)
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, (
            f"Operation should succeed with warning: {results[0].error_message if results else 'No results'}"
        )
        assert results[0].warning, "Operation should have warning flag (fallback to comment)"
        assert (
            "fallback" in results[0].error_message.lower()
            or "cross-cell" in results[0].error_message.lower()
        ), f"Expected fallback/cross-cell message, got: {results[0].error_message}"

        print("\n✓ Cell separator delete correctly fell back to comment")
        print(f"  Warning: {results[0].error_message}")
        print(f"  Output: {applier.output_path}")

    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_row_separator_delete_fallback(tmp_path):
    """
    Test that attempting to delete a row separator (\"], [\") falls back to comment.

    Scenario: Replace entire table with merged content that removes row separator
    Expected: Fallback to comment because deleting \"], [\" lands on boundary marker
    """
    print("\n" + "=" * 60)
    print("TEST: Row Separator Delete Fallback")
    print("=" * 60)

    # Create test document
    docx_path = tmp_path / 'test_row_separator_delete.docx'
    create_boundary_marker_table_document(docx_path)

    # Get paragraph IDs
    para_ids = _get_para_ids(docx_path)

    # Create JSONL that would merge rows (delete row separator)
    jsonl_path = tmp_path / 'test_row_separator_delete.jsonl'

    uuid_start = para_ids.get('Cell1')
    uuid_end = para_ids.get('Row2A')

    if not uuid_start or not uuid_end:
        print("ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"

    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')

        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '["Cell1", "Cell2", "Cell3"], ["Row2A", "Row2B", "Row2C"]',
            'violation_reason': 'Test row merge - should fallback to comment',
            'fix_action': 'replace',
            'revised_text': '["Cell1", "Cell2", "Cell3", "Row2A", "Row2B", "Row2C"]',
            'category': 'test',
            'rule_id': 'TEST_BOUNDARY_002',
            'heading': 'Test'
        }
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')

    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()

        # Check results - should succeed with warning (fallback to comment)
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, (
            f"Operation should succeed with warning: {results[0].error_message if results else 'No results'}"
        )
        assert results[0].warning, "Operation should have warning flag (fallback to comment)"
        assert (
            "fallback" in results[0].error_message.lower()
            or "cross-cell" in results[0].error_message.lower()
        ), f"Expected fallback/cross-cell message, got: {results[0].error_message}"

        print("\n✓ Row separator delete correctly fell back to comment")
        print(f"  Warning: {results[0].error_message}")
        print(f"  Output: {applier.output_path}")

    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def create_table_with_vmerge(para_ids):
    """
    Create a table with vertically merged cells.

    Structure:
      Col0    | Col1    | Col2
      --------+---------+--------
      Cell A  | Merged1 | Data1   (restart)
      Cell B  | Merged1 | Data2   (continue)
      Cell C  | Merged1 | Data3   (continue)

    Args:
        para_ids: List of paragraph IDs [p1, p2, p3, ...]

    Returns:
        Table element (w:tbl)
    """
    table_xml = f'''<w:tbl xmlns:w="{NS['w']}" xmlns:w14="{NS['w14']}">
        <w:tblGrid>
            <w:gridCol/>
            <w:gridCol/>
            <w:gridCol/>
        </w:tblGrid>
        <w:tr>
            <w:tc>
                <w:p w14:paraId="{para_ids[0]}">
                    <w:r><w:t>Cell A</w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:tcPr>
                    <w:vMerge w:val="restart"/>
                </w:tcPr>
                <w:p w14:paraId="{para_ids[1]}">
                    <w:r><w:t>Merged1</w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:p w14:paraId="{para_ids[2]}">
                    <w:r><w:t>Data1</w:t></w:r>
                </w:p>
            </w:tc>
        </w:tr>
        <w:tr>
            <w:tc>
                <w:p w14:paraId="{para_ids[3]}">
                    <w:r><w:t>Cell B</w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:tcPr>
                    <w:vMerge/>
                </w:tcPr>
                <w:p w14:paraId="{para_ids[4]}">
                    <w:r><w:t></w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:p w14:paraId="{para_ids[5]}">
                    <w:r><w:t>Data2</w:t></w:r>
                </w:p>
            </w:tc>
        </w:tr>
        <w:tr>
            <w:tc>
                <w:p w14:paraId="{para_ids[6]}">
                    <w:r><w:t>Cell C</w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:tcPr>
                    <w:vMerge/>
                </w:tcPr>
                <w:p w14:paraId="{para_ids[7]}">
                    <w:r><w:t></w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:p w14:paraId="{para_ids[8]}">
                    <w:r><w:t>Data3</w:t></w:r>
                </w:p>
            </w:tc>
        </w:tr>
    </w:tbl>'''

    return etree.fromstring(table_xml)


def create_table_with_gridspan(para_ids):
    """
    Create a table with horizontally merged cells (gridSpan).

    Structure:
      Col0         | Col1  | Col2
      -------------+-------+-------
      Wide Cell (spans 2)  | Data1
      Cell A       | Cell B| Data2

    Args:
        para_ids: List of paragraph IDs

    Returns:
        Table element (w:tbl)
    """
    table_xml = f'''<w:tbl xmlns:w="{NS['w']}" xmlns:w14="{NS['w14']}">
        <w:tblGrid>
            <w:gridCol/>
            <w:gridCol/>
            <w:gridCol/>
        </w:tblGrid>
        <w:tr>
            <w:tc>
                <w:tcPr>
                    <w:gridSpan w:val="2"/>
                </w:tcPr>
                <w:p w14:paraId="{para_ids[0]}">
                    <w:r><w:t>Wide Cell</w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:p w14:paraId="{para_ids[1]}">
                    <w:r><w:t>Data1</w:t></w:r>
                </w:p>
            </w:tc>
        </w:tr>
        <w:tr>
            <w:tc>
                <w:p w14:paraId="{para_ids[2]}">
                    <w:r><w:t>Cell A</w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:p w14:paraId="{para_ids[3]}">
                    <w:r><w:t>Cell B</w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:p w14:paraId="{para_ids[4]}">
                    <w:r><w:t>Data2</w:t></w:r>
                </w:p>
            </w:tc>
        </w:tr>
    </w:tbl>'''

    return etree.fromstring(table_xml)


def test_vmerge_restart_in_range():
    """Test vMerge restart within uuid range - content should be repeated"""
    print("Testing vMerge restart in range...")

    # Create a mock applier instance
    class MockApplier:
        def __init__(self):
            pass

        def _xpath(self, elem, expr):
            return elem.xpath(expr, namespaces=NS)

        def _find_ancestor_cell(self, para_elem):
            parent = para_elem.getparent()
            while parent is not None:
                if parent.tag == f'{{{NS["w"]}}}tc':
                    return parent
                parent = parent.getparent()
            return None

        def _find_ancestor_row(self, para_elem):
            parent = para_elem.getparent()
            while parent is not None:
                if parent.tag == f'{{{NS["w"]}}}tr':
                    return parent
                parent = parent.getparent()
            return None

        def _get_cell_merge_properties(self, tcPr):
            return AuditEditApplier._get_cell_merge_properties(self, tcPr)

        def _find_last_para_with_id_in_table(self, table_elem, uuid_end, start_para=None):
            return AuditEditApplier._find_last_para_with_id_in_table(
                self, table_elem, uuid_end, start_para
            )

        def _collect_runs_info_original(self, para_elem):
            # Simple text extraction for testing
            text = ''.join(t.text or '' for t in para_elem.findall('.//w:t', NS))
            if not text:
                return [], ''
            return [(para_elem, text, 0, len(text))], text

    applier = MockApplier()

    # Create table with vMerge
    para_ids = [f'{i:08X}' for i in range(9)]
    tbl = create_table_with_vmerge(para_ids)

    # Test: range includes restart (para_ids[1]) and continues (para_ids[4], [7])
    uuid_start = para_ids[1]  # Restart cell
    uuid_end = para_ids[7]    # Continue cell in last row

    # Call method under test
    data = AuditEditApplier._extract_text_in_range_from_table(
        applier, tbl, uuid_start, uuid_end
    )

    # Verify content - Merged1 should appear multiple times
    assert data
    assert "Merged1" in data
    assert data.count("Merged1") == 3, f"Expected Merged1 3 times, got: {data.count('Merged1')}"


def test_vmerge_restart_outside_range():
    """Test vMerge restart outside uuid range - content should not be repeated"""
    print("Testing vMerge restart outside range...")

    class MockApplier:
        def _xpath(self, elem, expr):
            return elem.xpath(expr, namespaces=NS)

        def _find_ancestor_cell(self, para_elem):
            parent = para_elem.getparent()
            while parent is not None:
                if parent.tag == f'{{{NS["w"]}}}tc':
                    return parent
                parent = parent.getparent()
            return None

        def _find_ancestor_row(self, para_elem):
            parent = para_elem.getparent()
            while parent is not None:
                if parent.tag == f'{{{NS["w"]}}}tr':
                    return parent
                parent = parent.getparent()
            return None

        def _get_cell_merge_properties(self, tcPr):
            return AuditEditApplier._get_cell_merge_properties(self, tcPr)

        def _find_last_para_with_id_in_table(self, table_elem, uuid_end, start_para=None):
            return AuditEditApplier._find_last_para_with_id_in_table(
                self, table_elem, uuid_end, start_para
            )

        def _collect_runs_info_original(self, para_elem):
            text = ''.join(t.text or '' for t in para_elem.findall('.//w:t', NS))
            if not text:
                return [], ''
            return [(para_elem, text, 0, len(text))], text

    applier = MockApplier()

    para_ids = [f'{i:08X}' for i in range(9)]
    tbl = create_table_with_vmerge(para_ids)

    # Range starts after the restart cell
    uuid_start = para_ids[2]
    uuid_end = para_ids[8]

    data = AuditEditApplier._extract_text_in_range_from_table(
        applier, tbl, uuid_start, uuid_end
    )

    assert data
    assert "Merged1" not in data, "Merged1 should not be included when restart is outside range"


def test_gridspan_cells_in_range():
    """Test gridSpan cells are repeated across columns within range."""
    print("Testing gridSpan in range...")

    class MockApplier:
        def _xpath(self, elem, expr):
            return elem.xpath(expr, namespaces=NS)

        def _find_ancestor_cell(self, para_elem):
            parent = para_elem.getparent()
            while parent is not None:
                if parent.tag == f'{{{NS["w"]}}}tc':
                    return parent
                parent = parent.getparent()
            return None

        def _find_ancestor_row(self, para_elem):
            parent = para_elem.getparent()
            while parent is not None:
                if parent.tag == f'{{{NS["w"]}}}tr':
                    return parent
                parent = parent.getparent()
            return None

        def _get_cell_merge_properties(self, tcPr):
            return AuditEditApplier._get_cell_merge_properties(self, tcPr)

        def _find_last_para_with_id_in_table(self, table_elem, uuid_end, start_para=None):
            return AuditEditApplier._find_last_para_with_id_in_table(
                self, table_elem, uuid_end, start_para
            )

        def _collect_runs_info_original(self, para_elem):
            text = ''.join(t.text or '' for t in para_elem.findall('.//w:t', NS))
            if not text:
                return [], ''
            return [(para_elem, text, 0, len(text))], text

    applier = MockApplier()

    para_ids = [f'{i:08X}' for i in range(5)]
    tbl = create_table_with_gridspan(para_ids)

    uuid_start = para_ids[0]
    uuid_end = para_ids[4]

    data = AuditEditApplier._extract_text_in_range_from_table(
        applier, tbl, uuid_start, uuid_end
    )

    assert data
    assert data.count("Wide Cell") == 2, f"Expected Wide Cell twice, got: {data.count('Wide Cell')}"


def _run_single_cell_extraction_flow() -> int:
    """Run the single-cell extraction flow and return 0 on success."""
    print("=" * 70)
    print("Test: Single-cell extraction from cross-cell match")
    print("=" * 70)

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        # 1. Create test document
        print("\n[Step 1] Creating test document...")
        docx_path = create_single_cell_test_document(temp_path)
        print(f"  Created: {docx_path}")

        # 2. Get paragraph IDs
        print("\n[Step 2] Extracting paragraph IDs...")
        para_ids = get_paragraph_ids(docx_path)
        print(f"  Found {len(para_ids)} paragraphs with IDs:")
        for i, pid in enumerate(para_ids):
            print(f"    [{i}] {pid}")

        if len(para_ids) < 2:
            print("\n  ERROR: Not enough paragraphs with IDs!")
            return 1

        # Use first paragraph as uuid_start, last paragraph as uuid_end
        uuid_start = para_ids[0]
        uuid_end = para_ids[-1]

        # 3. Create JSONL
        print("\n[Step 3] Creating JSONL with violation...")
        jsonl_path = Path(docx_path).with_suffix('.jsonl')
        create_single_cell_test_jsonl(docx_path, str(jsonl_path), uuid_start, uuid_end)
        print(f"  Created: {jsonl_path}")
        print(f"  UUID range: {uuid_start} -> {uuid_end}")

        # 4. Run apply_audit_edits with verbose mode
        print("\n[Step 4] Running apply_audit_edits with verbose mode...")
        print("-" * 70)

        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=False,
            verbose=True  # Enable verbose output to see debug logs
        )

        results = applier.apply()

        print("-" * 70)

        # 5. Check results
        print("\n[Step 5] Analyzing results...")
        for i, result in enumerate(results):
            print(f"\n  Result {i+1}:")
            print(f"    Success: {result.success}")
            print(f"    Warning: {result.warning if hasattr(result, 'warning') else 'N/A'}")
            print(f"    Error: {result.error_message if result.error_message else 'None'}")
            print(f"    Rule: {result.item.rule_id}")
            print(f"    Action: {result.item.fix_action}")

        # 6. Save output
        output_path = Path(docx_path).with_stem(Path(docx_path).stem + '_edited')
        applier.save()

        print(f"\n[Step 6] Output saved to: {output_path}")
        print("\n" + "=" * 70)
        print("Test complete!")
        print("=" * 70)

        return 0


def test_single_cell_extraction_main():
    """Run the single-cell extraction flow as a test."""
    assert _run_single_cell_extraction_flow() == 0


def create_single_cell_test_document(output_dir: Path) -> str:
    """Create a Word document with a simple table for testing."""
    doc = Document()

    # Add a paragraph before table (anchor paragraph)
    doc.add_paragraph("Table below:")

    # Add table with 1 row, 3 columns
    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0]
    row.cells[0].text = "72"
    row.cells[1].text = "军品电容"
    row.cells[2].text = "CAK55-D-10V-100uF-K"

    # Add a paragraph after table
    doc.add_paragraph("End of document")

    # Add w14:paraId attributes manually (required for apply_audit_edits.py)
    body_elem = doc._element.body
    para_counter = 0
    for para in body_elem.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        # Generate a unique paraId (8 hex characters)
        para_id = f"{para_counter:08X}"
        para.set('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', para_id)
        para_counter += 1

    # Save to temporary file in the provided directory
    docx_path = output_dir / "single_cell_extraction.docx"
    doc.save(docx_path)

    return str(docx_path)


def get_paragraph_ids(docx_path):
    """Extract paraId values from document."""
    doc = Document(docx_path)
    body_elem = doc._element.body

    para_ids = []
    for para in body_elem.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        para_id = para.get('{http://schemas.microsoft.com/office/word/2010/wordml}paraId')
        if para_id:
            para_ids.append(para_id)

    return para_ids


def create_single_cell_test_jsonl(docx_path, jsonl_path, uuid_start, uuid_end):
    """Create JSONL file with the test violation."""
    import hashlib

    # Calculate document hash
    sha256 = hashlib.sha256()
    with open(docx_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b""):
            sha256.update(chunk)
    doc_hash = f"sha256:{sha256.hexdigest()}"

    # Create JSONL content
    meta = {
        'type': 'meta',
        'source_file': docx_path,
        'source_hash': doc_hash
    }

    # Use raw strings to avoid JSON escaping - the script will handle escaping
    violation = {
        'category': 'consistency',
        'fix_action': 'replace',
        'violation_reason': 'Component name inconsistency: should be "军品钽电容"',
        'violation_text': '["72", "军品电容", "CAK55-D-10V-100uF-K"]',
        'revised_text': '["72", "军品钽电容", "CAK55-D-10V-100uF-K"]',
        'rule_id': 'R003',
        'uuid': uuid_start,
        'uuid_end': uuid_end,
        'heading': 'Test Heading'
    }

    with open(jsonl_path, 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False)
        f.write('\n')
        json.dump(violation, f, ensure_ascii=False)
        f.write('\n')


if __name__ == '__main__':
    sys.exit(_run_single_cell_extraction_flow())
