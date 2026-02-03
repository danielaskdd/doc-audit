#!/usr/bin/env python3
"""
ABOUTME: Unit tests for apply_audit_edits.py
"""

import sys
import json
import tempfile
from pathlib import Path

TESTS_DIR = Path(__file__).parent
sys.path.insert(0, str(TESTS_DIR))

import pytest
from lxml import etree
from unittest.mock import patch

import _apply_audit_edits_helpers as helpers
from _apply_audit_edits_helpers import (
    apply_module, AuditEditApplier, NS, DRAWING_PATTERN,
    strip_auto_numbering, EditItem, EditResult, NSMAP,
    create_paragraph_xml, create_paragraph_with_inline_image,
    create_paragraph_with_anchor_image, create_paragraph_with_track_changes,
    create_mock_applier, create_edit_item, get_test_author,
    create_mock_body_with_paragraphs, create_table_cell_with_paragraphs,
    create_table_with_cells, create_multi_row_table,
)


class TestRealDocumentTable:
    """Tests using the real test.docx file for table operations"""

    @staticmethod
    def get_test_doc_path():
        """Get path to test.docx"""
        return Path(__file__).parent / 'test.docx'

    def test_table_detection_in_real_document(self):
        """Test that paragraphs in test.docx table are correctly detected"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Find a paragraph in the table (Row 1, Col 2: "全自动贴片生产线")
        # ParaId: 04B34894
        para = applier._find_para_node_by_id('04B34894')
        if para is None:
            pytest.skip("ParaId 04B34894 not found in document")

        assert applier._is_paragraph_in_table(para) is True

        # Check that we can find the cell and row
        cell = applier._find_ancestor_cell(para)
        assert cell is not None

        row = applier._find_ancestor_row(para)
        assert row is not None

        table = applier._find_ancestor_table(para)
        assert table is not None

    def test_same_row_cells_no_row_boundary_error(self):
        """Test that cells in the same row don't trigger row boundary error"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 1: 751B37B6 (序号), 30353808 (项目), 04B34894 (设备工装需求), 6CF3F5AF (数量), 662522F3 (备注)
        # Use start: 30353808, end: 6CF3F5AF (same row)
        start_para = applier._find_para_node_by_id('30353808')
        if start_para is None:
            pytest.skip("ParaId 30353808 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '6CF3F5AF'
        )

        # Should not have boundary error for same row
        assert boundary_error is None
        assert is_cross_para is True  # Table mode is always cross-paragraph

    def test_different_rows_collected_with_row_marker(self):
        """Test that content from different rows is collected with row boundary marker"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 1 cell: 04B34894 (全自动贴片生产线)
        # Row 2 cell: 2B402507 (恒温干燥柜)
        start_para = applier._find_para_node_by_id('04B34894')
        if start_para is None:
            pytest.skip("ParaId 04B34894 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '2B402507'
        )

        # No upfront boundary error - content is collected across rows
        assert boundary_error is None

        # Content from both rows should be present
        assert '全自动贴片生产线' in combined_text
        assert '恒温干燥柜' in combined_text

        # Row boundary marker should be present
        assert '"], ["' in combined_text

        # _check_cross_row_boundary should detect the row span
        all_affected = applier._find_affected_runs(runs_info, 0, len(combined_text))
        real_runs = [r for r in all_affected
                     if not r.get('is_json_boundary', False)
                     and not r.get('is_json_escape', False)]
        assert applier._check_cross_row_boundary(real_runs) is True

    def test_json_format_for_real_table_row(self):
        """Test JSON format generation for a real table row"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 0 (header): 12C551B6 -> 5CC914F0
        # Expected: ["序号", "项目", "设备工装需求", "数量", "备注"]
        start_para = applier._find_para_node_by_id('12C551B6')
        if start_para is None:
            pytest.skip("ParaId 12C551B6 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '5CC914F0'
        )

        assert boundary_error is None
        # Check JSON format
        assert combined_text.startswith('["')
        assert combined_text.endswith('"]')
        assert '", "' in combined_text  # Cell separator

        # Verify content is present
        assert '序号' in combined_text
        assert '项目' in combined_text
        assert '设备工装需求' in combined_text
        assert '数量' in combined_text
        assert '备注' in combined_text

    def test_single_cell_in_real_table(self):
        """Test collecting runs from a single cell in real table"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Single cell: Row 1, Col 2 - "全自动贴片生产线" (paraId: 04B34894)
        start_para = applier._find_para_node_by_id('04B34894')
        if start_para is None:
            pytest.skip("ParaId 04B34894 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '04B34894'  # Same start and end
        )

        assert boundary_error is None
        # Single cell should produce ["content"]
        assert combined_text == '["全自动贴片生产线"]'

    def test_cross_cell_detection_in_real_table(self):
        """Test cross-cell boundary detection with real table"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 1: cells with paraIds 30353808, 04B34894, 6CF3F5AF
        # Collect runs spanning these cells
        start_para = applier._find_para_node_by_id('30353808')
        if start_para is None:
            pytest.skip("ParaId 30353808 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '6CF3F5AF'
        )

        assert boundary_error is None

        # Filter to get real runs (not JSON boundaries)
        real_runs = [r for r in runs_info
                     if not r.get('is_json_boundary', False)
                     and not r.get('is_json_escape', False)]

        # Should span multiple cells
        assert applier._check_cross_cell_boundary(real_runs) is True



class TestRealDocumentTableViolation:
    """Tests for processing violations in real document tables"""

    @staticmethod
    def get_test_doc_path():
        """Get path to test.docx"""
        return Path(__file__).parent / 'test.docx'

    def test_single_cell_violation_match(self):
        """Test matching violation text within a single cell"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 2, Col 2: "恒温干燥柜" (paraId: 2B402507)
        start_para = applier._find_para_node_by_id('2B402507')
        if start_para is None:
            pytest.skip("ParaId 2B402507 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '2B402507'
        )

        assert boundary_error is None

        # Search for "干燥" within the cell
        violation_text = '干燥'
        pos = combined_text.find(violation_text)

        # Should find it (accounting for JSON format ["恒温干燥柜"])
        assert pos != -1, f"'{violation_text}' not found in '{combined_text}'"

    def test_cross_cell_violation_in_same_row(self):
        """Test that violation spanning multiple cells in same row can be found"""
        from docx import Document as DocxDocument

        doc_path = self.get_test_doc_path()
        if not doc_path.exists():
            pytest.skip(f"Test file not found: {doc_path}")

        doc = DocxDocument(str(doc_path))
        applier = create_mock_applier()
        applier.body_elem = doc._element.body

        # Row 0 (header): all cells
        start_para = applier._find_para_node_by_id('12C551B6')
        if start_para is None:
            pytest.skip("ParaId 12C551B6 not found")

        runs_info, combined_text, is_cross_para, boundary_error = applier._collect_runs_info_across_paragraphs(
            start_para, '5CC914F0'
        )

        assert boundary_error is None

        # A cross-cell violation text would include the cell separator
        # For example: '项目", "设备工装需求'
        cross_cell_text = '项目", "设备工装需求'
        pos = combined_text.find(cross_cell_text)

        assert pos != -1, f"Cross-cell text '{cross_cell_text}' not found in '{combined_text}'"

        # Find affected runs
        match_end = pos + len(cross_cell_text)
        affected = applier._find_affected_runs(runs_info, pos, match_end)

        # Filter real runs
        real_runs = [r for r in affected
                     if not r.get('is_json_boundary', False)
                     and not r.get('is_json_escape', False)]

        # Should detect cross-cell boundary
        assert applier._check_cross_cell_boundary(real_runs) is True


# ============================================================
# Test Class: Filter Real Runs (Issue 1 - JSON boundary runs without elem/rPr)
# ============================================================
