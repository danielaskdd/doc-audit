#!/usr/bin/env python3
"""
Test that operations attempting to modify JSON boundary markers correctly fallback to comments.

This test verifies the fix for the issue where changes landing entirely on JSON boundary
markers (e.g., ", " cell separator or "], [" row separator) were silently ignored,
leading to partial application of edits.
"""

import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / 'skills' / 'doc-audit' / 'scripts'))

from lxml import etree
from docx import Document

from apply_audit_edits import AuditEditApplier, NS  # type: ignore


def create_test_table_document(docx_path: Path):
    """
    Create a test document with a table for boundary marker operations.
    
    Table structure:
    +-------+-------+-------+
    | Cell1 | Cell2 | Cell3 |
    +-------+-------+-------+
    | Row2A | Row2B | Row2C |
    +-------+-------+-------+
    """
    doc = Document()
    
    # Add title
    doc.add_heading('Boundary Marker Operations Test', 0)
    
    # Create table with 2 rows x 3 columns
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # Fill table data
    cells = [
        ['Cell1', 'Cell2', 'Cell3'],
        ['Row2A', 'Row2B', 'Row2C']
    ]
    
    for i, row_data in enumerate(cells):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    
    # Save document
    doc.save(str(docx_path))
    
    # Add paraId attributes (required for apply_audit_edits.py)
    _add_para_ids(docx_path)
    
    return doc


def _add_para_ids(docx_path: Path):
    """Add w14:paraId attributes to all paragraphs (simulate Word 2013+ behavior)."""
    doc = Document(str(docx_path))
    body_elem = doc._element.body
    
    # Add w14 namespace if not present
    nsmap = body_elem.nsmap
    if 'w14' not in nsmap:
        # Register namespace at document level
        etree.register_namespace('w14', NS['w14'])
    
    counter = 0
    for p in body_elem.iter(f'{{{NS["w"]}}}p'):
        # Generate unique 8-character hex ID (simulating Word's format)
        para_id = f'{counter:08X}'
        p.set('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', para_id)
        counter += 1
    
    # Save with updated paraId attributes
    doc.save(str(docx_path))


def get_para_ids(docx_path: Path) -> dict:
    """Extract paragraph IDs from the document for testing."""
    doc = Document(str(docx_path))
    body_elem = doc._element.body
    
    para_ids = {}
    for p in body_elem.iter(f'{{{NS["w"]}}}p'):
        para_id = p.get('{http://schemas.microsoft.com/office/word/2010/wordml}paraId')
        if para_id:
            # Extract text content
            text = ''.join(t.text or '' for t in p.iter(f'{{{NS["w"]}}}t'))
            para_ids[text] = para_id
    
    return para_ids


def test_cell_separator_delete_fallback():
    """
    Test that attempting to delete a cell separator (", ") falls back to comment.
    
    Scenario: Replace '["Cell1", "Cell2"]' with '["Cell1Cell2"]' (merge two cells)
    Expected: Fallback to comment because deleting ", " lands on boundary marker
    """
    print("\n" + "=" * 60)
    print("TEST: Cell Separator Delete Fallback")
    print("=" * 60)
    
    # Create test document
    test_dir = Path(__file__).parent
    docx_path = test_dir / 'test_cell_separator_delete.docx'
    create_test_table_document(docx_path)
    
    # Get paragraph IDs
    para_ids = get_para_ids(docx_path)
    
    # Create JSONL that would merge two cells (delete separator)
    jsonl_path = test_dir / 'test_cell_separator_delete.jsonl'
    
    uuid_start = para_ids.get('Cell1')
    uuid_end = para_ids.get('Cell2')
    
    if not uuid_start or not uuid_end:
        print(f"ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"
    
    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')
        
        # Replace operation that would delete the ", " separator
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '"Cell1", "Cell2"',
            'violation_reason': 'Test cell merge - should fallback to comment',
            'fix_action': 'replace',
            'revised_text': '"Cell1Cell2"',  # Merge - removes ", " separator
            'category': 'test',
            'rule_id': 'TEST_BOUNDARY_001',
            'heading': 'Test'
        }
        import json
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')
    
    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()
        
        # Check results - should succeed with warning (fallback to comment)
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, f"Operation should succeed with warning: {results[0].error_message if results else 'No results'}"
        assert results[0].warning, "Operation should have warning flag (fallback to comment)"
        assert "fallback" in results[0].error_message.lower() or "cross-cell" in results[0].error_message.lower(), \
            f"Expected fallback/cross-cell message, got: {results[0].error_message}"
        
        print("\n✓ Cell separator delete correctly fell back to comment")
        print(f"  Warning: {results[0].error_message}")
        print(f"  Output: {applier.output_path}")
            
    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_row_separator_delete_fallback():
    """
    Test that attempting to delete a row separator ("], [") falls back to comment.
    
    Scenario: Replace entire table with merged content that removes row separator
    Expected: Fallback to comment because deleting "], [" lands on boundary marker
    """
    print("\n" + "=" * 60)
    print("TEST: Row Separator Delete Fallback")
    print("=" * 60)
    
    # Create test document
    test_dir = Path(__file__).parent
    docx_path = test_dir / 'test_row_separator_delete.docx'
    create_test_table_document(docx_path)
    
    # Get paragraph IDs
    para_ids = get_para_ids(docx_path)
    
    # Create JSONL that would merge rows (delete row separator)
    jsonl_path = test_dir / 'test_row_separator_delete.jsonl'
    
    uuid_start = para_ids.get('Cell1')
    uuid_end = para_ids.get('Row2A')
    
    if not uuid_start or not uuid_end:
        print(f"ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"
    
    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')
        
        # Replace operation that would delete the "], [" row separator
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '["Cell1", "Cell2", "Cell3"], ["Row2A"',
            'violation_reason': 'Test row merge - should fallback to comment',
            'fix_action': 'replace',
            'revised_text': '["Cell1Cell2Cell3Row2A"',  # Merge rows - removes "], [" separator
            'category': 'test',
            'rule_id': 'TEST_BOUNDARY_002',
            'heading': 'Test'
        }
        import json
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')
    
    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()
        
        # Check results - should succeed with warning (fallback to comment)
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, f"Operation should succeed with warning: {results[0].error_message if results else 'No results'}"
        assert results[0].warning, "Operation should have warning flag (fallback to comment)"
        # Could be either "cross-row" or "cross-cell" depending on detection order
        assert any(keyword in results[0].error_message.lower() for keyword in ["fallback", "cross-cell", "cross-row"]), \
            f"Expected fallback/cross-cell/cross-row message, got: {results[0].error_message}"
        
        print("\n✓ Row separator delete correctly fell back to comment")
        print(f"  Warning: {results[0].error_message}")
        print(f"  Output: {applier.output_path}")
            
    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def main():
    """Run all boundary marker operation tests."""
    print("\n" + "=" * 60)
    print("BOUNDARY MARKER OPERATIONS TEST SUITE")
    print("=" * 60)
    
    tests = [
        ('Cell separator delete fallback', test_cell_separator_delete_fallback),
        ('Row separator delete fallback', test_row_separator_delete_fallback),
    ]
    
    results = []
    for name, test_func in tests:
        try:
            test_func()
            results.append((name, True))
        except Exception as e:
            print(f"\n✗ {name} crashed: {e}")
            import traceback
            traceback.print_exc()
            results.append((name, False))
    
    # Summary
    print("\n" + "=" * 60)
    print("TEST SUMMARY")
    print("=" * 60)
    
    passed = sum(1 for _, success in results if success)
    total = len(results)
    
    for name, success in results:
        status = "✓ PASS" if success else "✗ FAIL"
        print(f"{status}: {name}")
    
    print(f"\nPassed: {passed}/{total}")
    
    return 0 if passed == total else 1


if __name__ == '__main__':
    sys.exit(main())
