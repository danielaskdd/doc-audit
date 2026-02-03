#!/usr/bin/env python3
"""
Test multi-row and multi-cell delete/replace operations in apply_audit_edits.py
"""

import sys
import tempfile
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / 'skills' / 'doc-audit' / 'scripts'))

from lxml import etree
from docx import Document

from apply_audit_edits import AuditEditApplier, EditItem, NS


def create_test_table_document(docx_path: Path):
    """
    Create a test document with a table for multi-cell operations.
    
    Table structure:
    +-------+-------+-------+
    | Row 1 | Data1 | Data2 |
    +-------+-------+-------+
    | Row 2 | Data3 | Data4 |
    +-------+-------+-------+
    | Row 3 | Data5 | Data6 |
    +-------+-------+-------+
    """
    doc = Document()
    
    # Add title
    doc.add_heading('Multi-Cell Operations Test', 0)
    
    # Create table with 3 rows x 3 columns
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Fill table data
    cells = [
        ['Row 1', 'Data1', 'Data2'],
        ['Row 2', 'Data3', 'Data4'],
        ['Row 3', 'Data5', 'Data6']
    ]
    
    for i, row_data in enumerate(cells):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    
    # Save document
    doc.save(str(docx_path))
    
    # Add paraId attributes (required for apply_audit_edits.py)
    _add_para_ids(docx_path)
    
    return doc


def _add_para_ids(docx_path: Path):
    """Add w14:paraId attributes to all paragraphs (simulate Word 2013+ behavior)."""
    import random
    
    doc = Document(str(docx_path))
    body_elem = doc._element.body
    
    # Add w14 namespace if not present
    nsmap = body_elem.nsmap
    if 'w14' not in nsmap:
        # Register namespace at document level
        etree.register_namespace('w14', NS['w14'])
    
    counter = 0
    for p in body_elem.iter(f'{{{NS["w"]}}}p'):
        # Generate unique 8-character hex ID (simulating Word's format)
        para_id = f'{counter:08X}'
        p.set('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', para_id)
        counter += 1
    
    # Save with updated paraId attributes
    doc.save(str(docx_path))


def get_para_ids(docx_path: Path) -> dict:
    """Extract paragraph IDs from the document for testing."""
    doc = Document(str(docx_path))
    body_elem = doc._element.body
    
    para_ids = {}
    for p in body_elem.iter(f'{{{NS["w"]}}}p'):
        para_id = p.get('{http://schemas.microsoft.com/office/word/2010/wordml}paraId')
        if para_id:
            # Extract text content
            text = ''.join(t.text or '' for t in p.iter(f'{{{NS["w"]}}}t'))
            para_ids[text] = para_id
    
    return para_ids


def test_multi_row_delete(tmp_path):
    """Test deleting content across multiple table rows."""
    print("\n" + "=" * 60)
    print("TEST: Multi-row Delete")
    print("=" * 60)
    
    # Create test document
    docx_path = tmp_path / 'test_multi_row_delete.docx'
    create_test_table_document(docx_path)
    
    # Get paragraph IDs
    para_ids = get_para_ids(docx_path)
    
    # Create JSONL with multi-row delete (JSON format)
    # Delete first two rows: ["Row 1", "Data1", "Data2"], ["Row 2", "Data3", "Data4"]
    jsonl_path = tmp_path / 'test_multi_row_delete.jsonl'
    
    # Find start and end para IDs
    uuid_start = para_ids.get('Row 1')
    uuid_end = para_ids.get('Data4')
    
    if not uuid_start or not uuid_end:
        print(f"ERROR: Could not find paragraph IDs")
        print(f"Available IDs: {para_ids}")
        assert False, "Could not find required paragraph IDs"
    
    with open(jsonl_path, 'w', encoding='utf-8') as f:
        # Meta line (required)
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')
        
        # Edit item: delete multi-row content
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '["Row 1", "Data1", "Data2"], ["Row 2", "Data3", "Data4"]',
            'violation_reason': 'Multi-row delete test',
            'fix_action': 'delete',
            'revised_text': '',
            'category': 'test',
            'rule_id': 'TEST001',
            'heading': 'Test'
        }
        import json
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')
    
    # Apply edits
    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()
        
        # Check results
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, f"Multi-row delete failed: {results[0].error_message if results else 'No results'}"
        
        print("\n✓ Multi-row delete succeeded")
        print(f"  Output: {applier.output_path}")
            
    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_multi_cell_same_row_delete(tmp_path):
    """Test deleting content across multiple cells in the same row."""
    print("\n" + "=" * 60)
    print("TEST: Multi-cell Same Row Delete")
    print("=" * 60)
    
    # Create test document
    docx_path = tmp_path / 'test_multi_cell_delete.docx'
    create_test_table_document(docx_path)
    
    # Get paragraph IDs
    para_ids = get_para_ids(docx_path)
    
    # Delete cells in first row: ["Data1", "Data2"] (skipping "Row 1")
    jsonl_path = tmp_path / 'test_multi_cell_delete.jsonl'
    
    uuid_start = para_ids.get('Data1')
    uuid_end = para_ids.get('Data2')
    
    if not uuid_start or not uuid_end:
        print(f"ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"
    
    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')
        
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '"Data1", "Data2"',  # Same row, two cells
            'violation_reason': 'Multi-cell same row delete test',
            'fix_action': 'delete',
            'revised_text': '',
            'category': 'test',
            'rule_id': 'TEST002',
            'heading': 'Test'
        }
        import json
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')
    
    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()
        
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, f"Multi-cell same row delete failed: {results[0].error_message if results else 'No results'}"
        
        print("\n✓ Multi-cell same row delete succeeded")
        print(f"  Output: {applier.output_path}")
            
    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_single_cell_replace(tmp_path):
    """Test replace operation where all changes are in a single cell."""
    print("\n" + "=" * 60)
    print("TEST: Single-cell Replace (diff spans one cell)")
    print("=" * 60)
    
    # Create test document
    docx_path = tmp_path / 'test_single_cell_replace.docx'
    create_test_table_document(docx_path)
    
    # Get paragraph IDs
    para_ids = get_para_ids(docx_path)
    
    # Replace content in "Data1" cell only (even though violation_text spans multiple cells)
    jsonl_path = tmp_path / 'test_single_cell_replace.jsonl'
    
    uuid_start = para_ids.get('Data1')
    uuid_end = para_ids.get('Data2')
    
    if not uuid_start or not uuid_end:
        print(f"ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"
    
    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')
        
        # Violation text spans two cells, but changes are only in first cell
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '"Data1", "Data2"',
            'violation_reason': 'Single-cell replace test',
            'fix_action': 'replace',
            'revised_text': '"NewData", "Data2"',  # Only first cell changes
            'category': 'test',
            'rule_id': 'TEST003',
            'heading': 'Test'
        }
        import json
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')
    
    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()
        
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, f"Single-cell replace failed: {results[0].error_message if results else 'No results'}"
        
        print("\n✓ Single-cell replace succeeded")
        print(f"  Output: {applier.output_path}")
            
    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_multi_cell_replace_fallback(tmp_path):
    """Test replace operation where changes truly cross cell boundaries (should fallback to comment)."""
    print("\n" + "=" * 60)
    print("TEST: Multi-cell Replace Fallback (Cross-boundary)")
    print("=" * 60)
    
    # Create test document
    docx_path = tmp_path / 'test_multi_cell_replace_fallback.docx'
    create_test_table_document(docx_path)
    
    # Get paragraph IDs
    para_ids = get_para_ids(docx_path)
    
    jsonl_path = tmp_path / 'test_multi_cell_replace_fallback.jsonl'
    
    uuid_start = para_ids.get('Data1')
    uuid_end = para_ids.get('Data2')
    
    if not uuid_start or not uuid_end:
        print(f"ERROR: Could not find paragraph IDs")
        assert False, "Could not find required paragraph IDs"
    
    with open(jsonl_path, 'w', encoding='utf-8') as f:
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')
        
        # True cross-boundary change: merge two cells into one (Data1 + Data2 → Data1Data2)
        # This change crosses the cell boundary and should fallback to comment
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '"Data1", "Data2"',
            'violation_reason': 'Multi-cell replace fallback test - merge cells',
            'fix_action': 'replace',
            'revised_text': '"Data1Data2"',  # Merge both cells - crosses boundary
            'category': 'test',
            'rule_id': 'TEST004',
            'heading': 'Test'
        }
        import json
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')
    
    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()
        
        # Should succeed with warning (fallback to comment)
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, f"Multi-cell replace failed: {results[0].error_message if results else 'No results'}"
        assert results[0].warning, "Multi-cell replace should have fallen back to comment with warning"
        
        print("\n✓ Multi-cell replace correctly fell back to comment")
        print(f"  Warning: {results[0].error_message}")
        print(f"  Output: {applier.output_path}")
            
    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def test_multi_cell_distributed_replace(tmp_path):
    """Test replace operation with changes distributed across multiple cells (like g2/Hz → g²/Hz)."""
    print("\n" + "=" * 60)
    print("TEST: Multi-cell Distributed Replace (g2/Hz → g²/Hz pattern)")
    print("=" * 60)
    
    # Create test document
    docx_path = tmp_path / 'test_multi_cell_distributed_replace.docx'
    
    # Create custom table with pattern: Hz, g2/Hz, Hz, g2/Hz, Hz, g2/Hz
    doc = Document()
    doc.add_heading('Multi-Cell Distributed Replace Test', 0)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Fill with alternating pattern
    cells_content = ['Hz', 'g2/Hz', 'Hz', 'g2/Hz', 'Hz', 'g2/Hz']
    for i, content in enumerate(cells_content):
        table.rows[0].cells[i].text = content
    
    doc.save(str(docx_path))
    
    # Add paraId attributes (required for apply_audit_edits.py)
    import random
    doc = Document(str(docx_path))
    body_elem = doc._element.body
    
    # Register w14 namespace
    etree.register_namespace('w14', NS['w14'])
    
    counter = 0
    for p in body_elem.iter(f'{{{NS["w"]}}}p'):
        para_id = f'{counter:08X}'
        p.set('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', para_id)
        counter += 1
    
    doc.save(str(docx_path))
    
    # Get paragraph IDs in document order
    doc = Document(str(docx_path))
    body_elem = doc._element.body
    
    # Collect all paragraph IDs in order (don't use dict - repeated text will overwrite)
    all_para_data = []
    for p in body_elem.iter(f'{{{NS["w"]}}}p'):
        para_id = p.get('{http://schemas.microsoft.com/office/word/2010/wordml}paraId')
        if para_id:
            text = ''.join(t.text or '' for t in p.iter(f'{{{NS["w"]}}}t'))
            all_para_data.append((text, para_id))
    
    # Create JSONL with distributed replace (JSON format)
    jsonl_path = tmp_path / 'test_multi_cell_distributed_replace.jsonl'
    
    # Find first 'Hz' and last 'g2/Hz' explicitly
    uuid_start = None
    uuid_end = None
    
    for text, para_id in all_para_data:
        if text == 'Hz' and uuid_start is None:
            uuid_start = para_id  # First match
        if text == 'g2/Hz':
            uuid_end = para_id  # Keep updating to get last match
    
    if not uuid_start or not uuid_end:
        print(f"ERROR: Could not find paragraph IDs")
        print(f"Available IDs: {para_ids}")
        assert False, "Could not find required paragraph IDs"
    
    with open(jsonl_path, 'w', encoding='utf-8') as f:
        # Meta line (required)
        f.write('{"type": "meta", "source_file": "' + str(docx_path) + '", "source_hash": "skip"}\n')
        
        # Edit item: replace g2/Hz with g²/Hz in multiple cells
        edit_item = {
            'uuid': uuid_start,
            'uuid_end': uuid_end,
            'violation_text': '["Hz", "g2/Hz", "Hz", "g2/Hz", "Hz", "g2/Hz"]',
            'violation_reason': 'Multi-cell distributed replace test (like user report)',
            'fix_action': 'replace',
            'revised_text': '["Hz", "g²/Hz", "Hz", "g²/Hz", "Hz", "g²/Hz"]',  # Only g2/Hz changes
            'category': 'test',
            'rule_id': 'TEST005',
            'heading': 'Test'
        }
        import json
        f.write(json.dumps(edit_item, ensure_ascii=False) + '\n')
    
    # Apply edits
    try:
        applier = AuditEditApplier(
            str(jsonl_path),
            skip_hash=True,
            verbose=True
        )
        results = applier.apply()
        applier.save()
        
        # Check results - should succeed (not fallback to comment)
        assert len(results) == 1, f"Expected 1 result, got {len(results)}"
        assert results[0].success, f"Multi-cell distributed replace failed: {results[0].error_message if results else 'No results'}"
        assert not results[0].warning, f"Multi-cell distributed replace should not fallback to comment, but got warning: {results[0].error_message}"
        
        print("\n✓ Multi-cell distributed replace succeeded (track changes applied)")
        print(f"  Output: {applier.output_path}")
            
    except Exception as e:
        print(f"\n✗ Exception: {e}")
        import traceback
        traceback.print_exc()
        raise


def main():
    """Run all multi-cell operation tests."""
    print("\n" + "=" * 60)
    print("MULTI-CELL OPERATIONS TEST SUITE")
    print("=" * 60)
    
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        tests = [
            ('Multi-row delete', lambda: test_multi_row_delete(temp_path)),
            ('Multi-cell same row delete', lambda: test_multi_cell_same_row_delete(temp_path)),
            ('Single-cell replace', lambda: test_single_cell_replace(temp_path)),
            ('Multi-cell replace fallback', lambda: test_multi_cell_replace_fallback(temp_path)),
            ('Multi-cell distributed replace (g2/Hz → g²/Hz)', lambda: test_multi_cell_distributed_replace(temp_path)),
        ]
        
        results = []
        for name, test_func in tests:
            try:
                test_func()  # Will raise AssertionError if test fails
                results.append((name, True))  # No exception = success
            except Exception as e:
                print(f"\n✗ {name} crashed: {e}")
                import traceback
                traceback.print_exc()
                results.append((name, False))
    
    # Summary
    print("\n" + "=" * 60)
    print("TEST SUMMARY")
    print("=" * 60)
    
    passed = sum(1 for _, success in results if success)
    total = len(results)
    
    for name, success in results:
        status = "✓ PASS" if success else "✗ FAIL"
        print(f"{status}: {name}")
    
    print(f"\nPassed: {passed}/{total}")
    
    return 0 if passed == total else 1


if __name__ == '__main__':
    sys.exit(main())
