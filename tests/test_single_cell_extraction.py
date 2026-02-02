#!/usr/bin/env python3
"""
Test case to reproduce single-cell extraction failure.

This test replicates the user-reported issue where a replace operation
on a table row with changes in only one cell incorrectly falls back to
comment instead of applying track changes.

Test scenario:
- violation_text: ["72", "军品电容", "CAK55-D-10V-100uF-K"]
- revised_text: ["72", "军品钽电容", "CAK55-D-10V-100uF-K"]
- Only cell 2 changes: "军品电容" → "军品钽电容"
- Expected: Track change applied to cell 2
- Actual (bug): Fallback to comment with "Cross-cell track change not supported"
"""

import sys
import tempfile
from pathlib import Path
from docx import Document
import json

# Add skills/doc-audit/scripts to Python path
scripts_dir = Path(__file__).parent.parent / 'skills' / 'doc-audit' / 'scripts'
sys.path.insert(0, str(scripts_dir))

from apply_audit_edits import AuditEditApplier  # noqa: E402  # type: ignore


def create_test_document():
    """Create a Word document with a simple table for testing."""
    doc = Document()
    
    # Add a paragraph before table (anchor paragraph)
    doc.add_paragraph("Table below:")
    
    # Add table with 1 row, 3 columns
    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0]
    row.cells[0].text = "72"
    row.cells[1].text = "军品电容"
    row.cells[2].text = "CAK55-D-10V-100uF-K"
    
    # Add a paragraph after table
    doc.add_paragraph("End of document")
    
    # Add w14:paraId attributes manually (required for apply_audit_edits.py)
    body_elem = doc._element.body
    para_counter = 0
    for para in body_elem.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        # Generate a unique paraId (8 hex characters)
        para_id = f"{para_counter:08X}"
        para.set('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', para_id)
        para_counter += 1
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name


def get_paragraph_ids(docx_path):
    """Extract paraId values from document."""
    doc = Document(docx_path)
    body_elem = doc._element.body
    
    para_ids = []
    for para in body_elem.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        para_id = para.get('{http://schemas.microsoft.com/office/word/2010/wordml}paraId')
        if para_id:
            para_ids.append(para_id)
    
    return para_ids


def create_test_jsonl(docx_path, jsonl_path, uuid_start, uuid_end):
    """Create JSONL file with the test violation."""
    import hashlib
    
    # Calculate document hash
    sha256 = hashlib.sha256()
    with open(docx_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b""):
            sha256.update(chunk)
    doc_hash = f"sha256:{sha256.hexdigest()}"
    
    # Create JSONL content
    meta = {
        'type': 'meta',
        'source_file': docx_path,
        'source_hash': doc_hash
    }
    
    # Use raw strings to avoid JSON escaping - the script will handle escaping
    violation = {
        'category': 'consistency',
        'fix_action': 'replace',
        'violation_reason': 'Component name inconsistency: should be "军品钽电容"',
        'violation_text': '["72", "军品电容", "CAK55-D-10V-100uF-K"]',
        'revised_text': '["72", "军品钽电容", "CAK55-D-10V-100uF-K"]',
        'rule_id': 'R003',
        'uuid': uuid_start,
        'uuid_end': uuid_end,
        'heading': 'Test Heading'
    }
    
    with open(jsonl_path, 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False)
        f.write('\n')
        json.dump(violation, f, ensure_ascii=False)
        f.write('\n')


def main():
    """Run the test case."""
    print("=" * 70)
    print("Test: Single-cell extraction from cross-cell match")
    print("=" * 70)
    
    # 1. Create test document
    print("\n[Step 1] Creating test document...")
    docx_path = create_test_document()
    print(f"  Created: {docx_path}")
    
    # 2. Get paragraph IDs
    print("\n[Step 2] Extracting paragraph IDs...")
    para_ids = get_paragraph_ids(docx_path)
    print(f"  Found {len(para_ids)} paragraphs with IDs:")
    for i, pid in enumerate(para_ids):
        print(f"    [{i}] {pid}")
    
    if len(para_ids) < 2:
        print("\n  ERROR: Not enough paragraphs with IDs!")
        return 1
    
    # Use first paragraph as uuid_start, last paragraph as uuid_end
    uuid_start = para_ids[0]
    uuid_end = para_ids[-1]
    
    # 3. Create JSONL
    print(f"\n[Step 3] Creating JSONL with violation...")
    jsonl_path = Path(docx_path).with_suffix('.jsonl')
    create_test_jsonl(docx_path, str(jsonl_path), uuid_start, uuid_end)
    print(f"  Created: {jsonl_path}")
    print(f"  UUID range: {uuid_start} -> {uuid_end}")
    
    # 4. Run apply_audit_edits with verbose mode
    print(f"\n[Step 4] Running apply_audit_edits with verbose mode...")
    print("-" * 70)
    
    applier = AuditEditApplier(
        str(jsonl_path),
        skip_hash=False,
        verbose=True  # Enable verbose output to see debug logs
    )
    
    results = applier.apply()
    
    print("-" * 70)
    
    # 5. Check results
    print(f"\n[Step 5] Analyzing results...")
    for i, result in enumerate(results):
        print(f"\n  Result {i+1}:")
        print(f"    Success: {result.success}")
        print(f"    Warning: {result.warning if hasattr(result, 'warning') else 'N/A'}")
        print(f"    Error: {result.error_message if result.error_message else 'None'}")
        print(f"    Rule: {result.item.rule_id}")
        print(f"    Action: {result.item.fix_action}")
    
    # 6. Save output
    output_path = Path(docx_path).with_stem(Path(docx_path).stem + '_edited')
    applier.save()
    
    print(f"\n[Step 6] Output saved to: {output_path}")
    print("\n" + "=" * 70)
    print("Test complete!")
    print("=" * 70)
    
    # Cleanup note
    print(f"\nTest files saved for inspection:")
    print(f"  Source: {docx_path}")
    print(f"  JSONL: {jsonl_path}")
    print(f"  Output: {output_path}")
    print("\nReminder: These are temporary files in /tmp (or equivalent)")
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
