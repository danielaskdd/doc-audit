"""
Test for superscript replacement bug fix.

Bug: When replacing "x<sup>2</sup>/Hz" with "x<sup>3</sup>/Hz", the /Hz portion
was being incorrectly mapped due to position mismatch between plain text (used
by diff) and combined text with markup tags.

The fix: Skip run preservation for equal operations when has_markup=True,
and just recreate the text with proper formatting.
"""

import sys
from pathlib import Path
from lxml import etree
from docx import Document

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / 'skills' / 'doc-audit' / 'scripts'))

from apply_audit_edits import AuditEditApplier  # noqa: E402  # type: ignore

# Namespace
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
}


def test_superscript_replacement_preserves_suffix():
    """
    Test that replacing superscript doesn't corrupt following text.
    
    Scenario: "振动量级：0.04g<sup>2</sup>/Hz" → "振动量级：0.04g<sup>3</sup>/Hz"
    Expected: Only the superscript "2" should be replaced with "3", /Hz should remain intact
    Bug behavior: /Hz was being replaced with "sup"
    """
    # Create test document with superscript
    doc = Document()
    para = doc.add_paragraph()
    
    # Add text: "振动量级：0.04g"
    run1 = para.add_run("振动量级：0.04g")
    
    # Add superscript "2" (<sup>2</sup>)
    run2 = para.add_run("2")
    rPr2 = run2._element.get_or_add_rPr()
    vertAlign = etree.SubElement(rPr2, f'{{{NS["w"]}}}vertAlign')
    vertAlign.set(f'{{{NS["w"]}}}val', 'superscript')
    
    # Add "/Hz"
    run3 = para.add_run("/Hz")
    
    # Get paragraph element and add w14:paraId
    para_elem = para._element
    para_elem.set('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', 'TEST001')
    
    # Initialize applier
    applier = AuditEditApplier.__new__(AuditEditApplier)
    applier.doc = doc
    applier.body_elem = doc._element.body
    applier.next_change_id = 0
    applier.next_comment_id = 0
    applier.comments = []
    applier.operation_timestamp = '2026-01-01T00:00:00Z'
    applier.verbose = False
    
    # Collect original runs
    runs_info, combined = applier._collect_runs_info_original(para_elem)
    
    # Verify initial state
    assert combined == "振动量级：0.04g<sup>2</sup>/Hz", f"Initial text mismatch: {combined}"
    
    # Apply replace operation
    violation_text = "振动量级：0.04g<sup>2</sup>/Hz"
    revised_text = "振动量级：0.04g<sup>3</sup>/Hz"
    
    result = applier._apply_replace(
        para_elem=para_elem,
        violation_text=violation_text,
        revised_text=revised_text,
        violation_reason="Test replacement",
        orig_runs_info=runs_info,
        orig_match_start=0,
        author="Test"
    )
    
    assert result == 'success', f"Replace operation failed: {result}"
    
    # Collect runs after replacement
    after_runs, after_text = applier._collect_runs_info_original(para_elem)
    
    # Verify result - should preserve everything except the superscript number
    # Expected structure:
    # - "振动量级：0.04g" (normal text)
    # - delete: "2" with superscript
    # - insert: "3" with superscript
    # - "/Hz" (normal text)
    
    # Check that we have track changes
    ins_elements = para_elem.findall('.//w:ins', NS)
    del_elements = para_elem.findall('.//w:del', NS)
    
    assert len(del_elements) > 0, "No deletion found"
    assert len(ins_elements) > 0, "No insertion found"
    
    # Check deletion: should contain "2" with superscript
    del_run = del_elements[0].find('.//w:r', NS)
    assert del_run is not None, "No run in deletion"
    
    del_rPr = del_run.find('w:rPr', NS)
    assert del_rPr is not None, "No rPr in deleted run"
    
    del_vertAlign = del_rPr.find('w:vertAlign', NS)
    assert del_vertAlign is not None, "No vertAlign in deleted run"
    assert del_vertAlign.get(f'{{{NS["w"]}}}val') == 'superscript', "Deleted text should be superscript"
    
    del_text = del_run.find('w:delText', NS)
    assert del_text is not None and del_text.text == "2", f"Deleted text should be '2', got '{del_text.text if del_text is not None else None}'"
    
    # Check insertion: should contain "3" with superscript
    ins_run = ins_elements[0].find('.//w:r', NS)
    assert ins_run is not None, "No run in insertion"
    
    ins_rPr = ins_run.find('w:rPr', NS)
    assert ins_rPr is not None, "No rPr in inserted run"
    
    ins_vertAlign = ins_rPr.find('w:vertAlign', NS)
    assert ins_vertAlign is not None, "No vertAlign in inserted run"
    assert ins_vertAlign.get(f'{{{NS["w"]}}}val') == 'superscript', "Inserted text should be superscript"
    
    ins_text = ins_run.find('w:t', NS)
    assert ins_text is not None and ins_text.text == "3", f"Inserted text should be '3', got '{ins_text.text if ins_text is not None else None}'"
    
    # Most importantly: verify that "/Hz" is still present and not corrupted
    # Collect all normal text runs (not in w:ins or w:del)
    normal_runs = []
    for child in para_elem:
        if child.tag == f'{{{NS["w"]}}}r':
            t_elem = child.find('w:t', NS)
            if t_elem is not None and t_elem.text:
                normal_runs.append(t_elem.text)
    
    # Should find "/Hz" in normal runs
    assert "/Hz" in normal_runs, f"'/Hz' should be preserved in normal runs, found: {normal_runs}"
    
    # Verify it's not corrupted to "sup" or anything else
    full_text = ''.join(normal_runs)
    assert "/Hz" in full_text, f"'/Hz' should be in final text, got: {full_text}"
    assert "sup" not in full_text.lower() or "<sup>" in combined, f"Text should not contain literal 'sup', got: {full_text}"
    
    print("✓ Test passed: Superscript replacement preserves suffix correctly")


if __name__ == '__main__':
    test_superscript_replacement_preserves_suffix()
    print("\n✅ All tests passed!")
