#!/usr/bin/env python3
"""
Comprehensive tests for table_extractor.py

Tests include:
- Vertical merge (vMerge) content repetition
- Horizontal merge (gridSpan) handling
- ParaId extraction and consistency in merged cells
- Real Word document validation
"""

import sys
from pathlib import Path

# Add skills path
sys.path.insert(0, str(Path(__file__).parent.parent / 'skills' / 'doc-audit' / 'scripts'))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from table_extractor import TableExtractor  # type: ignore


# ============================================================================
# Helper Functions
# ============================================================================

def create_test_document_with_vmerge():
    """Create a test document with vertically merged cells with paraId attributes"""
    doc = Document()
    
    # Create a 3x3 table
    table = doc.add_table(rows=3, cols=3)
    
    # Add content to all cells first
    table.cell(0, 0).text = "Merged Content"
    table.cell(1, 0).text = "Row 2 content (to be replaced)"
    table.cell(2, 0).text = "Row 3 content (to be replaced)"
    
    table.cell(0, 1).text = "Normal Cell 1"
    table.cell(1, 1).text = "Normal Cell 2"
    table.cell(2, 1).text = "Normal Cell 3"
    
    table.cell(0, 2).text = "Column 3 Row 1"
    table.cell(1, 2).text = "Column 3 Row 2"
    table.cell(2, 2).text = "Column 3 Row 3"
    
    # Set vMerge on first column
    # Row 0, Col 0: restart
    tc_0_0 = table.cell(0, 0)._element
    tcPr_0_0 = tc_0_0.get_or_add_tcPr()
    vmerge_restart = OxmlElement('w:vMerge')
    vmerge_restart.set(qn('w:val'), 'restart')
    tcPr_0_0.append(vmerge_restart)
    
    # Row 1, Col 0: continue
    tc_1_0 = table.cell(1, 0)._element
    tcPr_1_0 = tc_1_0.get_or_add_tcPr()
    vmerge_continue = OxmlElement('w:vMerge')
    # No val attribute means continue
    tcPr_1_0.append(vmerge_continue)
    
    # Row 2, Col 0: continue
    tc_2_0 = table.cell(2, 0)._element
    tcPr_2_0 = tc_2_0.get_or_add_tcPr()
    vmerge_continue_2 = OxmlElement('w:vMerge')
    tcPr_2_0.append(vmerge_continue_2)
    
    # Add w14:paraId attributes to all paragraphs (simulate Word 2013+ behavior)
    body_elem = doc._element.body
    counter = 0
    for p in body_elem.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        para_id = f'{counter:08X}'
        p.set('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', para_id)
        counter += 1
    
    return doc, table


# ============================================================================
# Test Class: Vertical Merge with Synthetic Documents
# ============================================================================

class TestVMerge:
    """Test vertical merge behavior with programmatically created documents"""
    
    def test_vmerge_content_repetition(self):
        """Test that vertically merged cells repeat content in all rows"""
        doc, table = create_test_document_with_vmerge()
        result = TableExtractor.extract(table)
        
        # Check that all three rows in column 0 have the same content
        expected_content = "Merged Content"
        for i in range(3):
            actual = result[i][0]
            assert actual == expected_content, \
                f"Row {i}, Col 0: expected '{expected_content}', got '{actual}'"
        
        # Check that other columns have normal content
        assert result[0][1] == "Normal Cell 1"
        assert result[1][1] == "Normal Cell 2"
        assert result[2][1] == "Normal Cell 3"
    
    def test_vmerge_with_metadata(self):
        """Test that metadata extraction works with vertically merged cells"""
        doc, table = create_test_document_with_vmerge()
        result = TableExtractor.extract_with_metadata(table)
        
        # Verify structure
        assert 'rows' in result
        assert 'para_ids' in result
        assert 'para_ids_end' in result
        
        # Verify all merged rows have content
        rows = result['rows']
        expected_content = "Merged Content"
        for i in range(3):
            assert rows[i][0] == expected_content, \
                f"Row {i}, Col 0: expected '{expected_content}', got '{rows[i][0]}'"
    
    def test_vmerge_para_ids_end_uses_actual_para_id(self):
        """
        Test that vMerge continue cells use actual paraId for para_ids_end (range boundary).
        
        Design:
        - para_ids: All rows in vMerge use restart's paraId (for edit targeting)
        - para_ids_end: Each row uses its own actual paraId (for range boundary)
        """
        doc, table = create_test_document_with_vmerge()
        result = TableExtractor.extract_with_metadata(table)
        
        para_ids = result['para_ids']
        para_ids_end = result['para_ids_end']
        
        # Verify we have data for all 3 rows
        assert len(para_ids) == 3
        assert len(para_ids_end) == 3
        
        # Column 0 is vertically merged (restart in row 0, continue in rows 1-2)
        # para_ids should all be the same (restart's paraId) - for edit targeting
        restart_para_id = para_ids[0][0]
        assert restart_para_id is not None, "Restart cell should have paraId"
        assert para_ids[1][0] == restart_para_id, "Row 1 should share restart's paraId in para_ids"
        assert para_ids[2][0] == restart_para_id, "Row 2 should share restart's paraId in para_ids"
        
        # para_ids_end should be DIFFERENT (each row's actual paraId) - for range boundary
        # This is the key behavior after the fix
        para_id_end_0 = para_ids_end[0][0]
        para_id_end_1 = para_ids_end[1][0]
        para_id_end_2 = para_ids_end[2][0]
        
        assert para_id_end_0 is not None, "Row 0 should have para_id_end"
        assert para_id_end_1 is not None, "Row 1 should have para_id_end"
        assert para_id_end_2 is not None, "Row 2 should have para_id_end"
        
        # Each continue cell should use its own actual paraId for para_ids_end
        assert para_id_end_0 != para_id_end_1, \
            f"Row 0 and Row 1 should have different para_ids_end (got {para_id_end_0} vs {para_id_end_1})"
        assert para_id_end_1 != para_id_end_2, \
            f"Row 1 and Row 2 should have different para_ids_end (got {para_id_end_1} vs {para_id_end_2})"
        assert para_id_end_0 != para_id_end_2, \
            f"Row 0 and Row 2 should have different para_ids_end (got {para_id_end_0} vs {para_id_end_2})"


# ============================================================================
# Test Class: Real Word Document
# ============================================================================

class TestRealDocument:
    """Test extraction on real Word document (tests/test.docx)"""
    
    def test_real_document_extraction(self):
        """Test extraction on tests/test.docx"""
        doc_path = Path(__file__).parent / 'test.docx'
        
        assert doc_path.exists(), f"File not found: {doc_path}"
        
        doc = Document(str(doc_path))
        assert len(doc.tables) > 0, "No tables found in document"
        
        table = doc.tables[0]
        result = TableExtractor.extract_with_metadata(table)
        rows = result['rows']
        para_ids = result['para_ids']
        
        # Verify we have data
        assert len(rows) > 0
        assert len(para_ids) == len(rows)
    
    def test_vertical_merge_column_1(self):
        """Test vertical merge in '项目' column (index 1)"""
        doc_path = Path(__file__).parent / 'test.docx'
        doc = Document(str(doc_path))
        table = doc.tables[0]
        result = TableExtractor.extract_with_metadata(table)
        rows = result['rows']
        
        # Row 1-8 should have "PCBA电装生产设备"
        expected_text_1 = "PCBA电装生产设备"
        for i in range(1, 9):
            if i < len(rows) and len(rows[i]) > 1:
                assert expected_text_1 in rows[i][1], \
                    f"Row {i}, Col 1: expected to contain '{expected_text_1}', got '{rows[i][1][:50]}'"
        
        # Rows 9-16 should have "焊接质量检查"
        expected_text_2 = "焊接质量检查"
        for i in range(9, 17):
            if i < len(rows) and len(rows[i]) > 1:
                assert expected_text_2 in rows[i][1], \
                    f"Row {i}, Col 1: expected to contain '{expected_text_2}', got '{rows[i][1][:50]}'"
    
    def test_vertical_merge_column_4(self):
        """Test vertical merge in '备注' column (index 4)"""
        doc_path = Path(__file__).parent / 'test.docx'
        doc = Document(str(doc_path))
        table = doc.tables[0]
        result = TableExtractor.extract_with_metadata(table)
        rows = result['rows']
        
        # Rows 3-4 should have "焊膏存放、处理"
        expected_text = "焊膏存放、处理"
        for i in range(3, 5):
            if i < len(rows) and len(rows[i]) > 4:
                assert expected_text in rows[i][4], \
                    f"Row {i}, Col 4: expected to contain '{expected_text}', got '{rows[i][4]}'"
    
    def test_horizontal_merge_last_row(self):
        """Test horizontal merge in last row (row 17, '合计')"""
        doc_path = Path(__file__).parent / 'test.docx'
        doc = Document(str(doc_path))
        table = doc.tables[0]
        result = TableExtractor.extract_with_metadata(table)
        rows = result['rows']
        
        last_row_idx = 17
        assert last_row_idx < len(rows), f"Row {last_row_idx} not found (table has {len(rows)} rows)"
        
        last_row = rows[last_row_idx]
        assert len(last_row) > 1, "Last row doesn't have enough columns"
        assert "合计" in last_row[1], f"'合计' not found in Row {last_row_idx}, Col 1"
    
    def test_para_id_extraction(self):
        """Test that paraIds are correctly extracted"""
        doc_path = Path(__file__).parent / 'test.docx'
        doc = Document(str(doc_path))
        table = doc.tables[0]
        result = TableExtractor.extract_with_metadata(table)
        para_ids = result['para_ids']
        
        non_none_count = 0
        total_cells = 0
        
        for row in para_ids:
            for para_id in row:
                total_cells += 1
                if para_id is not None:
                    non_none_count += 1
        
        # Should have at least some paraIds from a real Word document
        assert non_none_count > 0, "No paraIds found (document may not have w14:paraId attributes)"
        
        percentage = (non_none_count / total_cells * 100) if total_cells > 0 else 0
        assert percentage > 50, f"Only {percentage:.1f}% of cells have paraIds (expected > 50%)"
    
    def test_vmerge_para_id_consistency(self):
        """Test that vertically merged cells share the same paraId"""
        doc_path = Path(__file__).parent / 'test.docx'
        doc = Document(str(doc_path))
        table = doc.tables[0]
        result = TableExtractor.extract_with_metadata(table)
        para_ids = result['para_ids']
        
        # Check "项目" column (Col 1): Row 1-8 should have same paraId
        if len(para_ids) > 8 and all(len(row) > 1 for row in para_ids[:9]):
            first_para_id = para_ids[1][1]  # Row 1, Col 1 (start of first merge)
            if first_para_id:
                for i in range(2, 9):
                    assert para_ids[i][1] == first_para_id, \
                        f"Row {i}, Col 1: paraId '{para_ids[i][1]}' != expected '{first_para_id}'"
        
        # Check "项目" column (Col 1): Row 9-16 should have same paraId
        if len(para_ids) > 16 and all(len(row) > 1 for row in para_ids[9:17]):
            second_para_id = para_ids[9][1]  # Row 9, Col 1 (start of second merge)
            if second_para_id:
                for i in range(10, 17):
                    assert para_ids[i][1] == second_para_id, \
                        f"Row {i}, Col 1: paraId '{para_ids[i][1]}' != expected '{second_para_id}'"
        
        # Check "备注" column (Col 4): Row 3-4 should have same paraId
        if len(para_ids) > 4 and all(len(row) > 4 for row in para_ids[3:5]):
            third_para_id = para_ids[3][4]  # Row 3, Col 4 (start of merge)
            if third_para_id:
                assert para_ids[4][4] == third_para_id, \
                    f"Row 4, Col 4: paraId '{para_ids[4][4]}' != expected '{third_para_id}'"
